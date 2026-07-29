"""Create the idempotent 100-case LegalBridge jury training workspace."""

from __future__ import annotations

import asyncio
import tempfile
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document
from sqlalchemy import func, select, update
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings, get_settings
from app.core.security import hash_password, verify_password
from app.db.base import utc_now
from app.db.session import Database
from app.models.analysis import AgentRun, AnalysisRun
from app.models.audit import AuditEvent
from app.models.auth_session import AuthSession
from app.models.case import LegalCase
from app.models.document import DocumentRecord
from app.models.document_page import DocumentPage
from app.models.enums import CaseStatus, UserRole
from app.models.organization import Organization
from app.models.user import User
from app.scripts.bootstrap_main import CaseSpec, SourceSpec, _ingest_source
from app.services.analysis import AGENTS, has_complete_flagship_data, run_case_analysis
from app.services.audit import add_audit_event
from app.services.storage import StorageService

ORGANIZATION_NAME = "LegalBridge Jury Legal-Aid Workspace"
ORGANIZATION_SLUG = "legalbridge-jury"
PRIMARY_EMAIL = "jury@legalbridge.local"
PRIMARY_PASSWORD = "LegalBridgeJury@2026"
PRIMARY_FULL_NAME = "LegalBridge Jury Presenter"
SAFETY_NOTICE = (
    "FICTIONAL TRAINING DATASET SAFETY NOTICE: This record contains no real "
    "person, client, allegation, official document, or legal conclusion. "
    "Not legal advice. Attorney review required. No automatic court filing."
)

CATEGORIES = (
    "Custodial procedure review",
    "Bail preparation",
    "Search and seizure review",
    "Witness inconsistency review",
    "Identification procedure review",
    "Property allegation defence",
    "Medical observation review",
    "Electronic evidence handling",
    "Juvenile procedure review",
    "Legal aid intake review",
)

COURTS = (
    "North District Legal Aid Review Forum",
    "Central District Pre-Trial Review Forum",
    "River Ward Magistrate Review Forum",
    "East District Youth Procedure Forum",
    "West District Evidence Review Forum",
)

JURISDICTIONS = (
    "Fictional North District jurisdiction",
    "Fictional Central District jurisdiction",
    "Fictional River Ward jurisdiction",
    "Fictional East District jurisdiction",
    "Fictional West District jurisdiction",
)

USER_SPECS = (
    (
        PRIMARY_EMAIL,
        PRIMARY_FULL_NAME,
        UserRole.ADMIN,
        PRIMARY_PASSWORD,
    ),
    (
        "attorney1@legalbridge.local",
        "Advocate Asha North",
        UserRole.ATTORNEY,
        "LegalBridgeAttorney1@2026",
    ),
    (
        "attorney2@legalbridge.local",
        "Advocate Kabir West",
        UserRole.ATTORNEY,
        "LegalBridgeAttorney2@2026",
    ),
    (
        "reviewer@legalbridge.local",
        "Senior Review Counsel Mira",
        UserRole.REVIEWER,
        "LegalBridgeReviewer@2026",
    ),
)


@dataclass(frozen=True)
class JuryBootstrapSummary:
    organization_id: str
    user_ids: dict[str, str]
    case_count: int
    completed_showcases: int
    pending_showcases: int
    supporting_cases: int
    document_count: int
    source_page_count: int
    audit_event_count: int
    best_completed_case: str
    best_pending_case: str


def _case_status(index: int) -> CaseStatus:
    if index <= 5:
        return CaseStatus.REVIEW
    if index <= 10:
        return CaseStatus.ACTIVE if index % 2 == 0 else CaseStatus.DRAFT
    return (
        CaseStatus.ACTIVE,
        CaseStatus.REVIEW,
        CaseStatus.DRAFT,
        CaseStatus.CLOSED,
    )[(index - 11) % 4]


def _case_spec(index: int) -> CaseSpec:
    category = CATEGORIES[(index - 1) % len(CATEGORIES)]
    title_suffix = (
        "Priority Source Review"
        if index <= 5
        else "Intake and Source Review"
        if index <= 10
        else f"Regional File {index:03d}"
    )
    return CaseSpec(
        number=f"LB-JURY-2026-{index:03d}",
        title=f"{category} — {title_suffix}",
        allegation_type=category,
        status=_case_status(index),
    )


def _document_count(index: int) -> int:
    if index <= 5:
        return 8
    if index <= 10:
        return 4
    if index <= 40:
        return 2
    return 1


def _source_spec(case_number: str, source_index: int) -> SourceSpec:
    source_types = (
        ("case-intake-record.pdf", "Case Intake Record", "intake record", "application/pdf"),
        (
            "procedure-observation.docx",
            "Procedure Observation",
            "procedure observation",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        ("event-chronology.txt", "Event Chronology", "chronology", "text/plain"),
        ("witness-account.pdf", "Witness Account", "witness account", "application/pdf"),
        (
            "property-inventory.docx",
            "Property Inventory",
            "property inventory",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        ("medical-note.txt", "Medical Observation Note", "medical observation", "text/plain"),
        (
            "identification-record.pdf",
            "Identification Procedure Record",
            "identification procedure",
            "application/pdf",
        ),
        (
            "electronic-record.docx",
            "Electronic Record Handling",
            "electronic record",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
    )
    suffix, title, category, content_type = source_types[source_index % len(source_types)]
    return SourceSpec(
        filename=f"{case_number.lower()}-{source_index + 1:02d}-{suffix}",
        title=title,
        category=category,
        content_type=content_type,
    )


def _source_paragraphs(case_spec: CaseSpec, source: SourceSpec) -> tuple[str, ...]:
    category = case_spec.allegation_type.lower()
    return (
        (
            f"File {case_spec.number} concerns a fictional {category}. "
            "Person A attended an intake meeting at 09:20 on 12 January 2026. "
            "The assigned legal-aid team recorded that the chronology remained provisional."
        ),
        (
            "Source comparison note: Record A describes an event at 10:15 near the "
            "north entrance, while Record B refers to 10:40 near the east corridor. "
            "The difference is an observation requiring attorney verification."
        ),
        (
            "Handling note: Item A was listed as a sealed paper packet. A later entry "
            "described a labelled document sleeve. The source sequence and custody "
            "description remain unresolved and should not be treated as a legal finding."
        ),
    )


def _write_pdf(path: Path, source: SourceSpec, case_spec: CaseSpec) -> None:
    document = pymupdf.open()
    try:
        for page_number, paragraph in enumerate(
            _source_paragraphs(case_spec, source), start=1
        ):
            page = document.new_page(width=595, height=842)
            page.insert_textbox(
                pymupdf.Rect(54, 54, 541, 788),
                (
                    f"{source.title}\n\nCase: {case_spec.number}\n"
                    f"Review category: {case_spec.allegation_type}\n"
                    f"Source section: {page_number}\n\n{paragraph}\n\n{SAFETY_NOTICE}"
                ),
                fontsize=11,
                fontname="helv",
                lineheight=1.35,
            )
        document.set_metadata(
            {
                "title": f"{case_spec.number} {source.title}",
                "author": "LegalBridge Jury Workspace",
                "subject": "Fictional synthetic legal-aid source",
            }
        )
        document.save(path, garbage=4, deflate=True)
    finally:
        document.close()


def _write_docx(path: Path, source: SourceSpec, case_spec: CaseSpec) -> None:
    document = Document()
    document.add_heading(source.title, level=0)
    document.add_paragraph(
        f"Case {case_spec.number} · {case_spec.allegation_type}"
    )
    for index, paragraph in enumerate(
        _source_paragraphs(case_spec, source), start=1
    ):
        document.add_heading(f"Source section {index}", level=1)
        document.add_paragraph(paragraph)
    document.add_heading("Safety notice", level=1)
    document.add_paragraph(SAFETY_NOTICE)
    document.core_properties.title = f"{case_spec.number} {source.title}"
    document.core_properties.subject = "Fictional synthetic legal-aid source"
    document.save(path)


def _write_txt(path: Path, source: SourceSpec, case_spec: CaseSpec) -> None:
    sections = [
        (
            f"{source.title.upper()} — SOURCE SECTION {index}\n"
            f"Case: {case_spec.number}\n"
            f"Review category: {case_spec.allegation_type}\n\n"
            f"{paragraph}\n\n{SAFETY_NOTICE}"
        )
        for index, paragraph in enumerate(
            _source_paragraphs(case_spec, source), start=1
        )
    ]
    path.write_text("\f".join(sections), encoding="utf-8")


def _generate_source(path: Path, source: SourceSpec, case_spec: CaseSpec) -> None:
    if path.suffix == ".pdf":
        _write_pdf(path, source, case_spec)
    elif path.suffix == ".docx":
        _write_docx(path, source, case_spec)
    else:
        _write_txt(path, source, case_spec)


async def _event_exists(
    session: AsyncSession,
    *,
    organization_id: str,
    event_type: str,
    entity_id: str,
) -> bool:
    event_id = await session.scalar(
        select(AuditEvent.id).where(
            AuditEvent.organization_id == organization_id,
            AuditEvent.event_type == event_type,
            AuditEvent.entity_id == entity_id,
        )
    )
    return event_id is not None


async def _provision_workspace(
    session: AsyncSession,
) -> tuple[Organization, list[User]]:
    organization = await session.scalar(
        select(Organization).where(Organization.slug == ORGANIZATION_SLUG)
    )
    if organization is None:
        organization = Organization(
            name=ORGANIZATION_NAME,
            slug=ORGANIZATION_SLUG,
            is_active=True,
        )
        session.add(organization)
        await session.flush()
    else:
        organization.name = ORGANIZATION_NAME
        organization.is_active = True

    users: list[User] = []
    for email, full_name, role, password in USER_SPECS:
        user = await session.scalar(
            select(User).where(
                User.organization_id == organization.id,
                User.email == email,
            )
        )
        if user is None:
            user = User(
                organization_id=organization.id,
                email=email,
                full_name=full_name,
                password_hash=hash_password(password),
                role=role,
                is_active=True,
            )
            session.add(user)
            await session.flush()
        else:
            user.full_name = full_name
            user.role = role
            user.is_active = True
            try:
                password_matches = verify_password(password, user.password_hash)
            except Exception:
                password_matches = False
            if not password_matches:
                user.password_hash = hash_password(password)
                user.token_version += 1
                await session.execute(
                    update(AuthSession)
                    .where(
                        AuthSession.user_id == user.id,
                        AuthSession.revoked_at.is_(None),
                    )
                    .values(revoked_at=utc_now(), updated_at=utc_now())
                )
        if not await _event_exists(
            session,
            organization_id=organization.id,
            event_type="jury_user_provisioned",
            entity_id=user.id,
        ):
            add_audit_event(
                session,
                organization_id=organization.id,
                actor_user_id=None,
                event_type="jury_user_provisioned",
                message="Jury workspace user provisioned for the synthetic legal-aid dataset.",
                entity_type="user",
                entity_id=user.id,
                metadata={"role": role.value, "synthetic": True},
            )
        users.append(user)
    if not await _event_exists(
        session,
        organization_id=organization.id,
        event_type="jury_workspace_provisioned",
        entity_id=organization.id,
    ):
        add_audit_event(
            session,
            organization_id=organization.id,
            actor_user_id=users[0].id,
            event_type="jury_workspace_provisioned",
            message="Jury legal-aid workspace provisioned.",
            entity_type="organization",
            entity_id=organization.id,
            metadata={"workspace_slug": ORGANIZATION_SLUG, "synthetic": True},
        )
    await session.commit()
    return organization, users


async def _provision_cases(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    attorneys: list[User],
) -> list[LegalCase]:
    cases: list[LegalCase] = []
    for index in range(1, 101):
        spec = _case_spec(index)
        legal_case = await session.scalar(
            select(LegalCase).where(
                LegalCase.organization_id == organization.id,
                LegalCase.case_number == spec.number,
            )
        )
        attorney = attorneys[(index - 1) % len(attorneys)]
        description = (
            f"{SAFETY_NOTICE} This matter supports a realistic source-linked "
            f"{spec.allegation_type.lower()} workflow."
        )
        if legal_case is None:
            legal_case = LegalCase(
                organization_id=organization.id,
                case_number=spec.number,
                title=spec.title,
                description=description,
                court_name=COURTS[(index - 1) % len(COURTS)],
                jurisdiction=JURISDICTIONS[(index - 1) % len(JURISDICTIONS)],
                allegation_type=spec.allegation_type,
                status=spec.status,
                created_by_id=primary.id,
                assigned_attorney_id=attorney.id,
            )
            session.add(legal_case)
            await session.flush()
        else:
            legal_case.title = spec.title
            legal_case.description = description
            legal_case.court_name = COURTS[(index - 1) % len(COURTS)]
            legal_case.jurisdiction = JURISDICTIONS[(index - 1) % len(JURISDICTIONS)]
            legal_case.allegation_type = spec.allegation_type
            legal_case.status = spec.status
            legal_case.assigned_attorney_id = attorney.id
        for event_type, message in (
            ("jury_case_created", "Legal-aid case record established."),
            ("jury_case_assigned", "Case assigned to fictional legal-aid counsel."),
            ("jury_case_status_recorded", "Case workflow status recorded."),
        ):
            if not await _event_exists(
                session,
                organization_id=organization.id,
                event_type=event_type,
                entity_id=legal_case.id,
            ):
                add_audit_event(
                    session,
                    organization_id=organization.id,
                    actor_user_id=primary.id,
                    event_type=event_type,
                    message=message,
                    entity_type="case",
                    entity_id=legal_case.id,
                    case_id=legal_case.id,
                    metadata={
                        "case_number": spec.number,
                        "status": spec.status.value,
                        "synthetic": True,
                    },
                )
        cases.append(legal_case)
    await session.commit()
    return cases


async def _provision_documents(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    cases: list[LegalCase],
    settings: Settings,
) -> None:
    storage = StorageService(settings.storage_root, settings.max_upload_bytes)
    storage.ensure_ready()
    with tempfile.TemporaryDirectory(prefix="legalbridge-jury-workspace-") as temporary:
        temporary_root = Path(temporary)
        for index, legal_case in enumerate(cases, start=1):
            case_spec = _case_spec(index)
            for source_index in range(_document_count(index)):
                source = _source_spec(case_spec.number, source_index)
                existing = await session.scalar(
                    select(DocumentRecord.id).where(
                        DocumentRecord.organization_id == organization.id,
                        DocumentRecord.case_id == legal_case.id,
                        DocumentRecord.original_filename == source.filename,
                    )
                )
                if existing is not None:
                    continue
                source_path = temporary_root / source.filename
                _generate_source(source_path, source, case_spec)
                await _ingest_source(
                    session,
                    source=source,
                    source_path=source_path,
                    organization=organization,
                    legal_case=legal_case,
                    primary=primary,
                    settings=settings,
                    storage=storage,
                )


async def _provision_pending_analysis(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    legal_case: LegalCase,
) -> None:
    existing = await session.scalar(
        select(AnalysisRun.id).where(
            AnalysisRun.organization_id == organization.id,
            AnalysisRun.case_id == legal_case.id,
            AnalysisRun.status.in_(("pending", "running")),
        )
    )
    if existing is not None:
        return
    now = utc_now()
    run = AnalysisRun(
        organization_id=organization.id,
        case_id=legal_case.id,
        status="running",
        provider="deterministic",
        started_by_user_id=primary.id,
        started_at=now,
        summary=(
            "Source extraction is complete. Deterministic analysis is in progress; "
            "no completed motion approval exists."
        ),
    )
    session.add(run)
    await session.flush()
    for sequence, (agent_key, agent_name) in enumerate(AGENTS, start=1):
        if sequence <= 3:
            status = "completed"
            output = "Initial source intake completed; attorney review remains pending."
            completed_at = now
        elif sequence == 4:
            status = "running"
            output = "Timeline review is in progress."
            completed_at = None
        else:
            status = "pending"
            output = "Waiting for preceding source-review stage."
            completed_at = None
        session.add(
            AgentRun(
                organization_id=organization.id,
                case_id=legal_case.id,
                analysis_run_id=run.id,
                agent_key=agent_key,
                agent_name=agent_name,
                sequence_number=sequence,
                status=status,
                input_summary="Extracted source pages available.",
                output_summary=output,
                started_at=now if sequence <= 4 else None,
                completed_at=completed_at,
            )
        )
    add_audit_event(
        session,
        organization_id=organization.id,
        actor_user_id=primary.id,
        event_type="jury_analysis_in_progress",
        message="Partial deterministic analysis state persisted.",
        entity_type="analysis_run",
        entity_id=run.id,
        case_id=legal_case.id,
        metadata={"completed_agents": 3, "synthetic": True},
    )
    await session.commit()


async def _provision_analysis(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    cases: list[LegalCase],
) -> None:
    for legal_case in cases[:5]:
        complete = await has_complete_flagship_data(
            session,
            organization_id=organization.id,
            case_id=legal_case.id,
        )
        if not complete:
            await run_case_analysis(
                session,
                organization_id=organization.id,
                case_id=legal_case.id,
                user_id=primary.id,
                provider_name="deterministic",
            )
    for legal_case in cases[5:10]:
        await _provision_pending_analysis(
            session,
            organization=organization,
            primary=primary,
            legal_case=legal_case,
        )


async def bootstrap_jury_workspace(
    database: Database,
    settings: Settings,
) -> JuryBootstrapSummary:
    if not settings.database_url.startswith("postgresql+asyncpg://"):
        raise RuntimeError("The jury workspace bootstrap requires active PostgreSQL.")
    async with database.session_factory() as session:
        organization, users = await _provision_workspace(session)
        primary = users[0]
        attorneys = [user for user in users if user.role == UserRole.ATTORNEY]
        cases = await _provision_cases(
            session,
            organization=organization,
            primary=primary,
            attorneys=attorneys,
        )
        await _provision_documents(
            session,
            organization=organization,
            primary=primary,
            cases=cases,
            settings=settings,
        )
        await _provision_analysis(
            session,
            organization=organization,
            primary=primary,
            cases=cases,
        )

        case_count = await session.scalar(
            select(func.count(LegalCase.id)).where(
                LegalCase.organization_id == organization.id
            )
        )
        document_count = await session.scalar(
            select(func.count(DocumentRecord.id)).where(
                DocumentRecord.organization_id == organization.id
            )
        )
        source_page_count = await session.scalar(
            select(func.count(DocumentPage.id)).where(
                DocumentPage.organization_id == organization.id
            )
        )
        audit_count = await session.scalar(
            select(func.count(AuditEvent.id)).where(
                AuditEvent.organization_id == organization.id
            )
        )
        completed_count = await session.scalar(
            select(func.count(AnalysisRun.id)).where(
                AnalysisRun.organization_id == organization.id,
                AnalysisRun.status == "completed",
            )
        )
        pending_count = await session.scalar(
            select(func.count(AnalysisRun.id)).where(
                AnalysisRun.organization_id == organization.id,
                AnalysisRun.status.in_(("pending", "running")),
            )
        )
        summary = JuryBootstrapSummary(
            organization_id=organization.id,
            user_ids={user.email: user.id for user in users},
            case_count=case_count or 0,
            completed_showcases=completed_count or 0,
            pending_showcases=pending_count or 0,
            supporting_cases=90,
            document_count=document_count or 0,
            source_page_count=source_page_count or 0,
            audit_event_count=audit_count or 0,
            best_completed_case="LB-JURY-2026-001",
            best_pending_case="LB-JURY-2026-006",
        )
        if (
            summary.case_count != 100
            or summary.completed_showcases < 5
            or summary.pending_showcases < 5
            or summary.document_count < 180
            or summary.source_page_count < 500
            or summary.audit_event_count < 500
        ):
            raise RuntimeError(f"Jury workspace minimums were not met: {summary}")
        return summary


async def main() -> None:
    settings = get_settings()
    database = Database(
        settings.database_url,
        echo=settings.sql_echo,
        ssl_mode=settings.database_ssl,
        pool_size=settings.database_pool_size,
        max_overflow=settings.database_max_overflow,
        pool_timeout=settings.database_pool_timeout,
        pool_recycle=settings.database_pool_recycle,
    )
    try:
        summary = await bootstrap_jury_workspace(database, settings)
    finally:
        await database.dispose()
    print(f"organization_id: {summary.organization_id}")
    for email, user_id in summary.user_ids.items():
        print(f"user_id[{email}]: {user_id}")
    print(f"total_cases: {summary.case_count}")
    print(f"completed_showcases: {summary.completed_showcases}")
    print(f"pending_showcases: {summary.pending_showcases}")
    print(f"supporting_cases: {summary.supporting_cases}")
    print(f"documents: {summary.document_count}")
    print(f"source_pages: {summary.source_page_count}")
    print(f"audit_events: {summary.audit_event_count}")
    print(f"best_completed_case: {summary.best_completed_case}")
    print(f"best_pending_case: {summary.best_pending_case}")


if __name__ == "__main__":
    asyncio.run(main())
