"""Create the professional ten-case LegalBridge casework workspace."""

from __future__ import annotations

import asyncio
import tempfile
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document
from docx.enum.text import WD_BREAK
from sqlalchemy import func, select

from app.core.config import Settings, get_settings
from app.core.security import hash_password, verify_password
from app.db.session import Database
from app.models.analysis import AnalysisRun, MotionDraft
from app.models.audit import AuditEvent
from app.models.case import LegalCase
from app.models.document import DocumentRecord
from app.models.document_page import DocumentPage
from app.models.enums import CaseStatus, UserRole
from app.models.organization import Organization
from app.models.user import User
from app.scripts.bootstrap_main import SourceSpec, _ingest_source
from app.services.analysis import run_case_analysis
from app.services.audit import add_audit_event
from app.services.storage import create_storage_service

ORGANIZATION_SLUG = "legalbridge-casework"
ORGANIZATION_NAME = "LegalBridge Casework"
FOOTER = "Fictional case record prepared for controlled platform evaluation."
REPOSITORY_ROOT = Path(__file__).resolve().parents[3]


@dataclass(frozen=True)
class UserSpec:
    email: str
    password: str
    role: UserRole
    full_name: str


@dataclass(frozen=True)
class CaseProfile:
    number: str
    title: str
    status: CaseStatus
    incident_date: str
    persons: tuple[str, str, str, str]
    locations: tuple[str, str]
    incident_times: tuple[str, str]
    custody_time: str
    arrest_time: str
    seal_codes: tuple[str, str]
    clothing: tuple[str, str]
    device: str
    export_times: tuple[str, str]
    medical_time: str
    transport_time: str


USERS = (
    UserSpec(
        "admin@legalbridge.in",
        "LegalBridgeAdmin@2026",
        UserRole.ADMIN,
        "LegalBridge Administrator",
    ),
    UserSpec(
        "attorney@legalbridge.in",
        "LegalBridgeAttorney@2026",
        UserRole.ATTORNEY,
        "Ananya Raman",
    ),
    UserSpec(
        "reviewer@legalbridge.in",
        "LegalBridgeReviewer@2026",
        UserRole.REVIEWER,
        "Karthik Menon",
    ),
)

CASE_TITLES = (
    "Custodial Timeline and Arrest Record Review",
    "Search and Seizure Chain-of-Custody Review",
    "Witness Identification and Description Review",
    "Electronic Record Timestamp and Integrity Review",
    "Bail Preparation and Medical Observation Review",
    "Property Recovery Chronology Review",
    "Witness Interview Sequence Review",
    "Missing CCTV Metadata Review",
    "Juvenile Procedure Intake Review",
    "Manual End-to-End Case",
)

NAMES = (
    ("Nikhil Sen", "Priya Nair", "Samar Iqbal", "Inspector Leena Thomas"),
    ("Rohan Iyer", "Kavya Das", "Imran Shah", "Inspector Vikram Bose"),
    ("Aditi Rao", "Farhan Ali", "Neel Joshi", "Inspector Tara Pillai"),
    ("Mohan Krish", "Sara Khan", "Deepak Roy", "Inspector Isha Varma"),
    ("Pooja Menon", "Arvind Lal", "Nitin Das", "Inspector Rhea Kapoor"),
    ("Kabir Anand", "Maya Sen", "Tarun Rao", "Inspector Diya Nair"),
    ("Leela Iyer", "Ritesh Paul", "Naveen Shah", "Inspector Om Prasad"),
    ("Sonia Das", "Harish Nair", "Yusuf Ali", "Inspector Mira Bose"),
    ("Ishan Rao", "Naina Sen", "Gopal Das", "Officer Reva Menon"),
    ("Arun Kumar", "Meera Rao", "Ravi Prakash", "Inspector Dev Menon"),
)

LOCATIONS = (
    ("Canal Gate", "Depot Lane entrance"),
    ("Market Arch", "Warehouse service yard"),
    ("Library Steps", "Community Hall passage"),
    ("Transit Plaza", "East Signal kiosk"),
    ("Clinic Junction", "River Road shelter"),
    ("Orchard Gate", "Mill Lane loading bay"),
    ("Civic Centre", "West Arcade"),
    ("Metro Walk", "South Parking exit"),
    ("Youth Centre", "Garden Road checkpoint"),
    ("North Gate", "Service Road Entrance"),
)


def _profiles() -> tuple[CaseProfile, ...]:
    profiles: list[CaseProfile] = []
    statuses = (
        CaseStatus.CLOSED,
        CaseStatus.CLOSED,
        CaseStatus.CLOSED,
        CaseStatus.CLOSED,
        CaseStatus.REVIEW,
        CaseStatus.REVIEW,
        CaseStatus.REVIEW,
        CaseStatus.ACTIVE,
        CaseStatus.ACTIVE,
        CaseStatus.DRAFT,
    )
    for index, title in enumerate(CASE_TITLES, start=1):
        is_manual = index == 10
        profiles.append(
            CaseProfile(
                number=f"LB-CASE-2026-{index:03d}",
                title=title,
                status=statuses[index - 1],
                incident_date="18 February 2026" if is_manual else f"{8 + index} February 2026",
                persons=NAMES[index - 1],
                locations=LOCATIONS[index - 1],
                incident_times=("20:15", "19:40")
                if is_manual
                else (f"{18 + index % 3}:15", f"{17 + index % 3}:40"),
                custody_time="20:50" if is_manual else f"{20 + index % 2}:05",
                arrest_time="21:30" if is_manual else f"{20 + index % 2}:45",
                seal_codes=("SP-18", "SP-81") if is_manual else (f"CR-{index}8", f"CR-8{index}"),
                clothing=("blue jacket", "grey shirt"),
                device=f"DEVICE-CW-{index:03d}",
                export_times=("22:10", "20:05")
                if is_manual
                else (f"{21 + index % 2}:10", f"{20 + index % 2}:55"),
                medical_time="22:35" if is_manual else f"{22 + index % 2}:05",
                transport_time="22:50" if is_manual else f"{22 + index % 2}:25",
            )
        )
    return tuple(profiles)


PROFILES = _profiles()

DOCUMENT_TYPES = (
    ("case_intake_note", "Case Intake Note", "intake"),
    ("witness_statement_a", "Witness Statement A", "witness_statement"),
    ("witness_statement_b", "Witness Statement B", "witness_statement"),
    ("arrest_memo", "Arrest and Custody Memo", "custody"),
    ("seizure_memo", "Seizure Memo and Property Inventory", "seizure"),
    ("chain_of_custody_log", "Chain-of-Custody Log", "evidence_handling"),
    ("electronic_record_log", "CCTV Export and Electronic Record Log", "electronic_record"),
    ("medical_observation", "Medical Observation and Transport Note", "medical"),
)


def _records(profile: CaseProfile, document_index: int) -> list[tuple[str, str, str]]:
    first, second, subject, officer = profile.persons
    location_a, location_b = profile.locations
    records = {
        0: [
            ("person", subject, f"{subject} attended intake with caseworker {first}."),
            (
                "incident_time",
                profile.incident_times[0],
                f"Initial account places the event near {location_a}.",
            ),
            ("location", location_a, "The intake sketch identifies the public entrance."),
        ],
        1: [
            ("person", first, f"{first} gave the first signed witness account."),
            (
                "incident_time",
                profile.incident_times[0],
                f"{first} observed the event after leaving the nearby kiosk.",
            ),
            ("location", location_a, f"{first} identified the approach beside {location_a}."),
            (
                "clothing",
                profile.clothing[0],
                f"{first} described the subject wearing a {profile.clothing[0]}.",
            ),
        ],
        2: [
            ("person", second, f"{second} gave a separately recorded account."),
            (
                "incident_time",
                profile.incident_times[1],
                f"{second} recalled the event before the street lights activated.",
            ),
            ("location", location_b, f"{second} identified the approach beside {location_b}."),
            ("clothing", profile.clothing[1], f"{second} described a {profile.clothing[1]}."),
        ],
        3: [
            ("person", officer, f"{officer} prepared the custody and arrest entries."),
            (
                "custody_time",
                profile.custody_time,
                f"The custody register records control beginning at {location_b}.",
            ),
            ("arrest_time", profile.arrest_time, "The arrest memo records the later formal time."),
            ("location", location_b, "The memo identifies the service-side entrance."),
        ],
        4: [
            (
                "seal_code",
                profile.seal_codes[0],
                "A mobile device was placed in a tamper-evident evidence pouch.",
            ),
            ("device_id", profile.device, "The device label was copied to the property inventory."),
            (
                "seizure_reference",
                profile.arrest_time,
                f"The seizure was recorded in the presence of {second}.",
            ),
        ],
        5: [
            (
                "seal_code",
                profile.seal_codes[1],
                "The receiving entry records a different seal code.",
            ),
            (
                "device_id",
                profile.device,
                "The same device identifier appears in the receiving register.",
            ),
            (
                "electronic_export_time",
                profile.export_times[1],
                "The station diary notes preparation for electronic export.",
            ),
        ],
        6: [
            (
                "electronic_export_time",
                profile.export_times[0],
                "The CCTV export log records completion of the copy.",
            ),
            (
                "electronic_certificate",
                "incomplete",
                "The certificate omits the hash-verification field and operator signature.",
            ),
            (
                "device_id",
                profile.device,
                "The export manifest links the copied file to the seized device.",
            ),
        ],
        7: [
            (
                "medical_time",
                profile.medical_time,
                f"The clinician recorded observations for {subject}.",
            ),
            (
                "transport_arrival_time",
                profile.transport_time,
                "The transport log records arrival after the observation time.",
            ),
            (
                "person",
                subject,
                "The note records the person as alert and able to answer questions.",
            ),
        ],
    }
    selected = records[document_index]
    if profile.number == "LB-CASE-2026-010" and document_index == 5:
        selected = [
            (
                "officers_departure_time",
                "20:05",
                "The station diary records officers leaving the station.",
            ),
            *selected,
        ]
    return selected


def _pages(profile: CaseProfile, document_index: int) -> list[str]:
    slug, title, _ = DOCUMENT_TYPES[document_index]
    officer = profile.persons[3]
    records = _records(profile, document_index)
    pages: list[str] = []
    for page_number in range(1, 4):
        selected = records[(page_number - 1) :: 3] or [records[(page_number - 1) % len(records)]]
        record_text = "\n".join(
            f"RECORD: {key} | {value} | {detail}" for key, value, detail in selected
        )
        page_focus = (
            "This page records the initial account and identifies the people "
            "responsible for the entry."
            if page_number == 1
            else "This page cross-references the related chronology and preserves "
            "the exact recorded value."
            if page_number == 2
            else "This page records evidence handling, unresolved differences, "
            "and the next source to compare."
        )
        pages.append(
            f"{title}\n"
            f"Case number: {profile.number}\n"
            f"Document date: {profile.incident_date}\n"
            f"Author or recording officer: {officer}\n"
            f"Record section: {page_number} of 3\n\n"
            f"{page_focus}\n\n{record_text}\n\n"
            f"Related record: {slug.replace('_', ' ').title()} entry {page_number}; "
            f"compare with the case chronology where values differ.\n\n{FOOTER}"
        )
    return pages


def _write_pdf(path: Path, title: str, pages: list[str]) -> None:
    document = pymupdf.open()
    for page_text in pages:
        page = document.new_page(width=595, height=842)
        page.insert_textbox(
            pymupdf.Rect(54, 58, 541, 790),
            page_text,
            fontsize=10.5,
            lineheight=1.35,
            fontname="helv",
        )
    document.set_metadata({"title": title, "subject": "Realistic fictional case record"})
    document.save(path)
    document.close()


def _write_docx(path: Path, title: str, pages: list[str]) -> None:
    document = Document()
    document.core_properties.title = title
    document.core_properties.subject = "Realistic fictional case record"
    for index, page_text in enumerate(pages):
        lines = page_text.splitlines()
        document.add_heading(lines[0], level=1)
        for line in lines[1:]:
            document.add_paragraph(line)
        if index < len(pages) - 1:
            document.paragraphs[-1].add_run().add_break(WD_BREAK.PAGE)
    document.save(path)


def _write_txt(path: Path, pages: list[str]) -> None:
    path.write_text("\n\f\n".join(pages), encoding="utf-8")


def _source(profile: CaseProfile, document_index: int) -> SourceSpec:
    slug, title, category = DOCUMENT_TYPES[document_index]
    extensions = (
        (".pdf", "application/pdf"),
        (".docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
        (".txt", "text/plain"),
    )
    extension, content_type = extensions[document_index % len(extensions)]
    return SourceSpec(
        filename=f"{document_index + 1:02d}_{slug}{extension}",
        title=title,
        category=category,
        content_type=content_type,
    )


def _write_source(path: Path, source: SourceSpec, profile: CaseProfile, index: int) -> None:
    pages = _pages(profile, index)
    if source.filename.endswith(".pdf"):
        _write_pdf(path, source.title, pages)
    elif source.filename.endswith(".docx"):
        _write_docx(path, source.title, pages)
    else:
        _write_txt(path, pages)


async def _organization_and_users(session) -> tuple[Organization, dict[str, User]]:
    organization = await session.scalar(
        select(Organization).where(Organization.slug == ORGANIZATION_SLUG)
    )
    if organization is None:
        organization = Organization(name=ORGANIZATION_NAME, slug=ORGANIZATION_SLUG, is_active=True)
        session.add(organization)
        await session.flush()
    users: dict[str, User] = {}
    for spec in USERS:
        user = await session.scalar(
            select(User).where(User.organization_id == organization.id, User.email == spec.email)
        )
        if user is None:
            user = User(
                organization_id=organization.id,
                email=spec.email,
                password_hash=hash_password(spec.password),
                full_name=spec.full_name,
                role=spec.role,
                is_active=True,
            )
            session.add(user)
            await session.flush()
        else:
            if not verify_password(spec.password, user.password_hash):
                user.password_hash = hash_password(spec.password)
            user.full_name = spec.full_name
            user.role = spec.role
            user.is_active = True
        users[spec.email] = user
    await session.commit()
    return organization, users


async def _cases(session, organization: Organization, users: dict[str, User]) -> list[LegalCase]:
    admin = users["admin@legalbridge.in"]
    attorney = users["attorney@legalbridge.in"]
    cases: list[LegalCase] = []
    for profile in PROFILES:
        legal_case = await session.scalar(
            select(LegalCase).where(
                LegalCase.organization_id == organization.id,
                LegalCase.case_number == profile.number,
            )
        )
        if legal_case is None:
            legal_case = LegalCase(
                organization_id=organization.id,
                case_number=profile.number,
                title=profile.title,
                description=(
                    f"Realistic fictional case records concerning {profile.title.lower()}, "
                    "with source-linked review questions and no real personal data."
                ),
                court_name=f"{profile.locations[0]} District Review Forum",
                jurisdiction="Fictional Indian district jurisdiction",
                allegation_type=profile.title.replace(" Review", ""),
                status=profile.status,
                created_by_id=admin.id,
                assigned_attorney_id=attorney.id,
            )
            session.add(legal_case)
            await session.flush()
            add_audit_event(
                session,
                organization_id=organization.id,
                actor_user_id=admin.id,
                event_type="case_created",
                message=f"{profile.number} created with complete professional metadata.",
                entity_type="case",
                entity_id=legal_case.id,
                case_id=legal_case.id,
            )
        else:
            legal_case.title = profile.title
            legal_case.status = profile.status
            legal_case.assigned_attorney_id = attorney.id
        cases.append(legal_case)
    await session.commit()
    return cases


def create_manual_case_pack(repository_root: Path) -> Path:
    profile = PROFILES[9]
    destination = repository_root / "case-packs" / profile.number
    destination.mkdir(parents=True, exist_ok=True)
    requested = (
        ("01_case_intake_summary.pdf", 0),
        ("02_witness_statement_arun_kumar.docx", 1),
        ("03_witness_statement_meera_rao.docx", 2),
        ("04_arrest_memo.pdf", 3),
        ("05_seizure_record.pdf", 4),
        ("06_medical_observation.docx", 7),
        ("07_electronic_record_log.txt", 6),
        ("08_station_diary_extract.txt", 5),
    )
    for filename, document_index in requested:
        path = destination / filename
        pages = _pages(profile, document_index)
        if filename == "08_station_diary_extract.txt":
            pages = [
                page.replace("Chain-of-Custody Log", "Station Diary Extract", 1) for page in pages
            ]
        if path.suffix == ".pdf":
            _write_pdf(path, DOCUMENT_TYPES[document_index][1], pages)
        elif path.suffix == ".docx":
            _write_docx(path, DOCUMENT_TYPES[document_index][1], pages)
        else:
            _write_txt(path, pages)
    instructions = """# Upload instructions for LB-CASE-2026-010

1. Sign in to the `legalbridge-casework` workspace.
2. Open `LB-CASE-2026-010`.
3. Upload all eight files in this folder.
4. Confirm that every document reaches **Processed** status.
5. Open **Extracted sources** and inspect the extracted pages.
6. Run analysis.
7. Check the timeline.
8. Check contradictions.
9. Check the procedural audit.
10. Ask Copilot questions about witness times, locations, seal codes,
    the electronic certificate, and medical chronology.
11. Generate a motion.
12. Run citation and ethics checks, then submit the motion for review.
13. Approve the reviewed version using the configured review PIN.
14. Export both PDF and DOCX.
"""
    (destination / "UPLOAD_INSTRUCTIONS.md").write_text(instructions, encoding="utf-8")
    return destination


async def bootstrap_casework_workspace(database: Database, settings: Settings) -> None:
    async with database.session_factory() as session:
        organization, users = await _organization_and_users(session)
        cases = await _cases(session, organization, users)
        admin = users["admin@legalbridge.in"]
        storage = create_storage_service(settings)
        storage.ensure_ready()
        with tempfile.TemporaryDirectory(prefix="legalbridge-casework-") as temporary:
            root = Path(temporary)
            for case_index, legal_case in enumerate(cases[:9]):
                profile = PROFILES[case_index]
                document_count = 8 if case_index < 4 else 6 if case_index < 7 else 3
                for document_index in range(document_count):
                    source = _source(profile, document_index)
                    source_path = root / f"{profile.number}-{source.filename}"
                    _write_source(source_path, source, profile, document_index)
                    await _ingest_source(
                        session,
                        source=source,
                        source_path=source_path,
                        organization=organization,
                        legal_case=legal_case,
                        primary=admin,
                        settings=settings,
                        storage=storage,
                    )

        for case_index, legal_case in enumerate(cases[:7]):
            existing = await session.scalar(
                select(AnalysisRun.id).where(
                    AnalysisRun.organization_id == organization.id,
                    AnalysisRun.case_id == legal_case.id,
                    AnalysisRun.status == "completed",
                )
            )
            if existing is None:
                await run_case_analysis(
                    session,
                    organization_id=organization.id,
                    case_id=legal_case.id,
                    user_id=admin.id,
                    provider_name=settings.analysis_provider,
                )
            if case_index >= 4:
                motions = (
                    await session.scalars(
                        select(MotionDraft).where(
                            MotionDraft.organization_id == organization.id,
                            MotionDraft.case_id == legal_case.id,
                        )
                    )
                ).all()
                for motion in motions:
                    motion.status = "submitted_for_review"
                await session.commit()

        pack = create_manual_case_pack(REPOSITORY_ROOT)
        case_count = await session.scalar(
            select(func.count(LegalCase.id)).where(LegalCase.organization_id == organization.id)
        )
        document_count = await session.scalar(
            select(func.count(DocumentRecord.id)).where(
                DocumentRecord.organization_id == organization.id
            )
        )
        page_count = await session.scalar(
            select(func.count(DocumentPage.id)).where(
                DocumentPage.organization_id == organization.id
            )
        )
        audit_count = await session.scalar(
            select(func.count(AuditEvent.id)).where(AuditEvent.organization_id == organization.id)
        )
        print(f"organization_id: {organization.id}")
        for email, user in users.items():
            print(f"user_id[{email}]: {user.id}")
        print(f"cases: {case_count}")
        print(f"documents: {document_count}")
        print(f"source_pages: {page_count}")
        print(f"audit_events: {audit_count}")
        print(f"manual_case_pack: {pack}")


async def main() -> None:
    settings = get_settings()
    database = Database(
        settings.database_url,
        echo=settings.sql_echo,
        ssl_mode=settings.database_ssl,
        pool_size=settings.database_pool_size,
        max_overflow=settings.database_max_overflow,
        pool_timeout=settings.database_pool_timeout,
        pool_recycle=settings.database_pool_recycle,
    )
    try:
        await bootstrap_casework_workspace(database, settings)
    finally:
        await database.dispose()


if __name__ == "__main__":
    asyncio.run(main())
