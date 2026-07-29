"""Generate, store, extract, and persist three synthetic demonstration sources."""

from __future__ import annotations

import asyncio
import tempfile
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings, get_settings
from app.db.base import new_uuid, utc_now
from app.db.session import Database
from app.models.audit import AuditEvent
from app.models.case import LegalCase
from app.models.document import DocumentRecord
from app.models.organization import Organization
from app.models.user import User
from app.scripts.bootstrap_demo import (
    DEMO_ADMIN_EMAIL,
    DEMO_CASE_NUMBER,
    DEMO_ORGANIZATION_SLUG,
)
from app.services.audit import add_audit_event
from app.services.document_processing import (
    document_audit_metadata,
    process_document,
)
from app.services.storage import StorageService


@dataclass(frozen=True)
class DemoDocument:
    filename: str
    category: str
    content_type: str


DEMO_DOCUMENTS = (
    DemoDocument(
        filename="synthetic-court-transcript.pdf",
        category="court transcript",
        content_type="application/pdf",
    ),
    DemoDocument(
        filename="synthetic-police-report.docx",
        category="police report",
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    DemoDocument(
        filename="synthetic-arrest-memo.txt",
        category="arrest memo",
        content_type="text/plain",
    ),
)

SYNTHETIC_NOTICE = (
    "SYNTHETIC HACKATHON DATA — fictional names and identifiers. "
    "Not an official court or police record. No real personal data."
)


def _write_pdf(path: Path) -> None:
    pages = [
        (
            "Synthetic Court Transcript — Page 1\n\n"
            f"{SYNTHETIC_NOTICE}\n\n"
            "Matter: LB-DEMO-2026-001\n"
            "Fictional witness: Asha Rao\n"
            "The witness states that the property file was shown at 09:40 on 12 June 2026.\n\n"
            "This transcript records fictional statements only and makes no legal conclusion."
        ),
        (
            "Synthetic Court Transcript — Page 2\n\n"
            f"{SYNTHETIC_NOTICE}\n\n"
            "Fictional witness: Dev Sen\n"
            "The witness recalls seeing a sealed envelope at 10:15 but did not "
            "inspect its contents.\n\n"
            "Counsel notes that time references require independent attorney verification."
        ),
        (
            "Synthetic Court Transcript — Page 3\n\n"
            f"{SYNTHETIC_NOTICE}\n\n"
            "Demonstration clerk note\n"
            "The fictional hearing adjourned at 11:05. No finding of fact or law was made.\n\n"
            "End of synthetic transcript."
        ),
    ]
    document = pymupdf.open()
    try:
        for content in pages:
            page = document.new_page(width=595, height=842)
            page.insert_textbox(
                pymupdf.Rect(54, 54, 541, 788),
                content,
                fontsize=11,
                fontname="helv",
                lineheight=1.35,
            )
        document.set_metadata(
            {
                "title": "Synthetic LegalBridge Court Transcript",
                "author": "LegalBridge India Hackathon",
                "subject": "Synthetic data; not an official record",
            }
        )
        document.save(path, garbage=4, deflate=True)
    finally:
        document.close()


def _write_docx(path: Path) -> None:
    document = Document()
    document.add_heading("Synthetic Police Report", level=0)
    document.add_paragraph(SYNTHETIC_NOTICE)
    document.add_heading("Demonstration identifiers", level=1)
    identifiers = document.add_table(rows=1, cols=2)
    identifiers.style = "Table Grid"
    identifiers.rows[0].cells[0].text = "Field"
    identifiers.rows[0].cells[1].text = "Synthetic value"
    for field, value in (
        ("Case", "LB-DEMO-2026-001"),
        ("Report", "FIC-PR-2026-014"),
        ("Reporting officer", "Inspector Mira Das (fictional)"),
    ):
        cells = identifiers.add_row().cells
        cells[0].text = field
        cells[1].text = value
    document.add_heading("Narrative", level=1)
    document.add_paragraph(
        "At 09:55 on 12 June 2026, the fictional reporting officer received a sealed "
        "property-paper envelope. The officer did not express a legal conclusion."
    )
    document.add_heading("Items recorded", level=1)
    items = document.add_table(rows=1, cols=3)
    items.style = "Table Grid"
    for index, heading in enumerate(("Item", "Fictional identifier", "Condition")):
        items.rows[0].cells[index].text = heading
    for values in (
        ("Envelope", "FIC-ENV-01", "Sealed"),
        ("Photocopy set", "FIC-COPY-02", "Three synthetic pages"),
    ):
        cells = items.add_row().cells
        for index, value in enumerate(values):
            cells[index].text = value
    document.add_heading("Review limitation", level=1)
    document.add_paragraph(
        "This generated report is synthetic source material for extraction demonstration. "
        "It does not establish an offence, violation, or admissible fact."
    )
    document.core_properties.title = "Synthetic LegalBridge Police Report"
    document.core_properties.subject = "Synthetic data; not an official record"
    document.save(path)


def _write_txt(path: Path) -> None:
    sections = [
        (
            "SYNTHETIC ARREST MEMO — SECTION 1\n"
            f"{SYNTHETIC_NOTICE}\n\n"
            "Case: LB-DEMO-2026-001\n"
            "Memo: FIC-AM-2026-008\n"
            "Fictional person: Rohan Mehta\n"
        ),
        (
            "SYNTHETIC ARREST MEMO — SECTION 2\n"
            "A demonstration entry records an announced time of 08:50 on 12 June 2026.\n"
            "A separate fictional station entry records 09:10.\n"
            "The difference is source text only and requires attorney verification.\n"
        ),
        (
            "SYNTHETIC ARREST MEMO — SECTION 3\n"
            "No unsupported legal conclusion is asserted.\n"
            "This file is not an official arrest memo and contains no real personal data.\n"
        ),
    ]
    path.write_text("\f".join(sections), encoding="utf-8")


def generate_demo_files(directory: Path) -> dict[str, Path]:
    paths = {definition.filename: directory / definition.filename for definition in DEMO_DOCUMENTS}
    _write_pdf(paths["synthetic-court-transcript.pdf"])
    _write_docx(paths["synthetic-police-report.docx"])
    _write_txt(paths["synthetic-arrest-memo.txt"])
    return paths


async def _ingest_demo_file(
    session: AsyncSession,
    *,
    definition: DemoDocument,
    source_path: Path,
    organization: Organization,
    legal_case: LegalCase,
    admin: User,
    settings: Settings,
    storage: StorageService,
) -> DocumentRecord:
    existing = await session.scalar(
        select(DocumentRecord).where(
            DocumentRecord.organization_id == organization.id,
            DocumentRecord.case_id == legal_case.id,
            DocumentRecord.original_filename == definition.filename,
        )
    )
    if existing is not None:
        return existing

    document_id = new_uuid()
    add_audit_event(
        session,
        organization_id=organization.id,
        actor_user_id=admin.id,
        event_type="document_upload_started",
        message="Synthetic demonstration document ingestion started.",
        entity_type="document_record",
        entity_id=document_id,
        case_id=legal_case.id,
        metadata={
            "document_id": document_id,
            "category": definition.category,
            "declared_content_type": definition.content_type,
        },
    )
    await session.commit()
    staged = storage.stage_file(
        source_path,
        filename=definition.filename,
        content_type=definition.content_type,
    )
    storage_key: str | None = None
    try:
        duplicate = await session.scalar(
            select(DocumentRecord.id).where(
                DocumentRecord.case_id == legal_case.id,
                DocumentRecord.sha256 == staged.sha256,
            )
        )
        if duplicate is not None:
            return await session.get(DocumentRecord, duplicate)
        storage_key = storage.finalize(
            staged,
            organization_id=organization.id,
            case_id=legal_case.id,
            document_id=document_id,
        )
        record = DocumentRecord(
            id=document_id,
            organization_id=organization.id,
            case_id=legal_case.id,
            original_filename=definition.filename,
            content_type=definition.content_type,
            size_bytes=staged.size_bytes,
            sha256=staged.sha256,
            category=definition.category,
            status="metadata_only",
            storage_key=storage_key,
            storage_backend="local_private",
            extraction_status="uploaded",
            original_uploaded_at=utc_now(),
            created_by_id=admin.id,
        )
        session.add(record)
        await session.flush()
        add_audit_event(
            session,
            organization_id=organization.id,
            actor_user_id=admin.id,
            event_type="document_uploaded",
            message="Synthetic demonstration document stored privately.",
            entity_type="document_record",
            entity_id=record.id,
            case_id=legal_case.id,
            metadata=document_audit_metadata(record),
        )
        await session.commit()
    except Exception:
        await session.rollback()
        if storage_key:
            storage.delete_key(storage_key)
        raise
    finally:
        storage.discard(staged)

    await process_document(
        session,
        record=record,
        storage=storage,
        settings=settings,
        actor_user_id=admin.id,
    )
    return record


async def bootstrap_demo_documents(database: Database, settings: Settings) -> int:
    storage = StorageService(settings.storage_root, settings.max_upload_bytes)
    storage.ensure_ready()
    with tempfile.TemporaryDirectory(prefix="legalbridge-demo-documents-") as temporary:
        generated = generate_demo_files(Path(temporary))
        async with database.session_factory() as session:
            organization = await session.scalar(
                select(Organization).where(Organization.slug == DEMO_ORGANIZATION_SLUG)
            )
            if organization is None:
                raise RuntimeError("Run the base demo bootstrap before document bootstrap.")
            legal_case = await session.scalar(
                select(LegalCase).where(
                    LegalCase.organization_id == organization.id,
                    LegalCase.case_number == DEMO_CASE_NUMBER,
                )
            )
            admin = await session.scalar(
                select(User).where(
                    User.organization_id == organization.id,
                    User.email == DEMO_ADMIN_EMAIL,
                )
            )
            if legal_case is None or admin is None:
                raise RuntimeError("The demonstration case or administrator is missing.")

            records = []
            for definition in DEMO_DOCUMENTS:
                records.append(
                    await _ingest_demo_file(
                        session,
                        definition=definition,
                        source_path=generated[definition.filename],
                        organization=organization,
                        legal_case=legal_case,
                        admin=admin,
                        settings=settings,
                        storage=storage,
                    )
                )

            bootstrap_event = await session.scalar(
                select(AuditEvent.id).where(
                    AuditEvent.organization_id == organization.id,
                    AuditEvent.event_type == "demo_documents_bootstrapped",
                    AuditEvent.entity_id == legal_case.id,
                )
            )
            if bootstrap_event is None:
                add_audit_event(
                    session,
                    organization_id=organization.id,
                    actor_user_id=admin.id,
                    event_type="demo_documents_bootstrapped",
                    message=(
                        "Three stored and extracted synthetic demonstration documents are ready."
                    ),
                    entity_type="case",
                    entity_id=legal_case.id,
                    case_id=legal_case.id,
                    metadata={
                        "document_count": len(records),
                        "filenames": [record.original_filename for record in records],
                    },
                )
                await session.commit()
            return len(records)


async def main() -> None:
    settings = get_settings()
    database = Database(
        settings.database_url,
        echo=settings.sql_echo,
        ssl_mode=settings.database_ssl,
        pool_size=settings.database_pool_size,
        max_overflow=settings.database_max_overflow,
        pool_timeout=settings.database_pool_timeout,
        pool_recycle=settings.database_pool_recycle,
    )
    try:
        count = await bootstrap_demo_documents(database, settings)
    finally:
        await database.dispose()
    print(f"LegalBridge synthetic source documents are ready ({count} records).")


if __name__ == "__main__":
    asyncio.run(main())
