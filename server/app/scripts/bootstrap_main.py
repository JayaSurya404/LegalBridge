"""Seed the hosted LegalBridge jury workspace with synthetic database-backed data."""

from __future__ import annotations

import asyncio
import secrets
import tempfile
from dataclasses import dataclass
from pathlib import Path

import pymupdf
from docx import Document
from sqlalchemy import func, select, update
from sqlalchemy.engine import make_url
from sqlalchemy.ext.asyncio import AsyncSession

from app.core.config import Settings, get_settings
from app.core.security import hash_password, verify_password
from app.db.base import new_uuid, utc_now
from app.db.session import Database
from app.models.audit import AuditEvent
from app.models.auth_session import AuthSession
from app.models.case import LegalCase
from app.models.document import DocumentRecord
from app.models.document_page import DocumentPage
from app.models.enums import CaseStatus, UserRole
from app.models.organization import Organization
from app.models.user import User
from app.services.audit import add_audit_event
from app.services.document_processing import document_audit_metadata, process_document
from app.services.storage import StorageService

MAIN_ORGANIZATION_NAME = "LegalBridge Main Jury Workspace"
MAIN_ORGANIZATION_SLUG = "legalbridge-main"
PRIMARY_FULL_NAME = "LegalBridge Main Demonstration Admin"
PRIMARY_EMAIL = "legalbridge@legalbridge.demo"
PRIMARY_PASSWORD = "legalbridge@2026"
SYNTHETIC_NOTICE = (
    "FICTIONAL SYNTHETIC HACKATHON DEMONSTRATION MATERIAL. "
    "Not an official legal, court, police, medical, or government record. "
    "Contains no real personal information."
)


@dataclass(frozen=True)
class StaffSpec:
    full_name: str
    email: str
    role: UserRole


@dataclass(frozen=True)
class CaseSpec:
    number: str
    title: str
    allegation_type: str
    status: CaseStatus


@dataclass(frozen=True)
class SourceSpec:
    filename: str
    title: str
    category: str
    content_type: str


@dataclass(frozen=True)
class BootstrapSummary:
    organization_id: str
    primary_user_id: str
    supporting_user_count: int
    case_count: int
    document_count: int
    source_page_count: int
    audit_event_count: int
    flagship_document_count: int


STAFF = (
    StaffSpec(
        "Adv. Ananya Rao (Synthetic)",
        "ananya.rao@legalbridge.demo",
        UserRole.ATTORNEY,
    ),
    StaffSpec(
        "Adv. Kabir Sen (Synthetic)",
        "kabir.sen@legalbridge.demo",
        UserRole.ATTORNEY,
    ),
    StaffSpec(
        "Justice Review Officer Mira Das (Synthetic)",
        "mira.das@legalbridge.demo",
        UserRole.REVIEWER,
    ),
    StaffSpec(
        "LegalBridge Jury Operations Admin (Synthetic)",
        "jury.admin@legalbridge.demo",
        UserRole.ADMIN,
    ),
)

CASE_TOPICS = (
    ("Comprehensive Synthetic Defence Demonstration", "Property allegation"),
    ("Synthetic Custodial Procedure Review", "Custodial procedure"),
    ("Synthetic Bail Preparation File", "Bail preparation"),
    ("Synthetic Search and Seizure Review", "Search and seizure"),
    ("Synthetic Witness Consistency Matter", "Witness inconsistency"),
    ("Synthetic Identification Procedure Review", "Identification procedure"),
    ("Synthetic Juvenile Safeguard File", "Juvenile procedure"),
    ("Synthetic Legal-Aid Intake Review", "Legal-aid intake"),
    ("Synthetic Medical Record Review", "Medical-record review"),
    ("Synthetic Electronic Record Handling", "Electronic-record handling"),
    ("Synthetic Arrest Procedure Review", "Arrest-procedure review"),
    ("Synthetic Seizure Record Review", "Seizure-record review"),
    ("Synthetic Property Registry Comparison", "Property allegation"),
    ("Synthetic Bail Compliance Archive", "Bail preparation"),
    ("Synthetic Legal-Aid Closure File", "Legal-aid intake"),
)
CASE_STATUSES = (
    CaseStatus.ACTIVE,
    CaseStatus.ACTIVE,
    CaseStatus.ACTIVE,
    CaseStatus.ACTIVE,
    CaseStatus.ACTIVE,
    CaseStatus.REVIEW,
    CaseStatus.REVIEW,
    CaseStatus.REVIEW,
    CaseStatus.DRAFT,
    CaseStatus.DRAFT,
    CaseStatus.DRAFT,
    CaseStatus.CLOSED,
    CaseStatus.CLOSED,
    CaseStatus.ARCHIVED,
    CaseStatus.ARCHIVED,
)
CASES = tuple(
    CaseSpec(
        number=f"LB-MAIN-2026-{index:03d}",
        title=title,
        allegation_type=allegation_type,
        status=CASE_STATUSES[index - 1],
    )
    for index, (title, allegation_type) in enumerate(CASE_TOPICS, start=1)
)

FLAGSHIP_SOURCES = (
    SourceSpec(
        "synthetic-multi-page-court-transcript.pdf",
        "Multi-page Court Transcript",
        "court transcript",
        "application/pdf",
    ),
    SourceSpec(
        "synthetic-police-report.docx",
        "Police Report",
        "police report",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    SourceSpec(
        "synthetic-arrest-memo.txt",
        "Arrest Memo",
        "arrest memo",
        "text/plain",
    ),
    SourceSpec(
        "synthetic-witness-statement.pdf",
        "Witness Statement",
        "witness statement",
        "application/pdf",
    ),
    SourceSpec(
        "synthetic-seizure-record.docx",
        "Seizure Record",
        "seizure record",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
    SourceSpec(
        "synthetic-medical-observation.txt",
        "Medical Observation",
        "medical observation",
        "text/plain",
    ),
    SourceSpec(
        "synthetic-identification-procedure.pdf",
        "Identification Procedure Record",
        "identification procedure",
        "application/pdf",
    ),
    SourceSpec(
        "synthetic-electronic-evidence-inventory.docx",
        "Electronic Evidence Inventory",
        "electronic evidence",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ),
)


def _supporting_sources(case_number: str) -> tuple[SourceSpec, ...]:
    stem = case_number.lower()
    return (
        SourceSpec(
            f"{stem}-source-summary.pdf",
            "Case Source Summary",
            "source summary",
            "application/pdf",
        ),
        SourceSpec(
            f"{stem}-procedure-record.docx",
            "Procedure Record",
            "procedure record",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ),
        SourceSpec(
            f"{stem}-review-note.txt",
            "Review Note",
            "review note",
            "text/plain",
        ),
    )


def _write_pdf(path: Path, source: SourceSpec, case: CaseSpec) -> None:
    document = pymupdf.open()
    try:
        for page_number in range(1, 4):
            page = document.new_page(width=595, height=842)
            body = (
                f"{source.title} - Synthetic Page {page_number}\n\n"
                f"{SYNTHETIC_NOTICE}\n\n"
                f"Case: {case.number}\n"
                f"Matter: {case.title}\n"
                f"Category: {source.category}\n\n"
                f"Controlled observation {page_number}: this fictional source records "
                "a demonstration timestamp, handling step, and reviewer note solely to "
                "exercise server-side validation and extraction.\n\n"
                "No allegation, violation, authority, or legal conclusion is established."
            )
            page.insert_textbox(
                pymupdf.Rect(54, 54, 541, 788),
                body,
                fontsize=11,
                fontname="helv",
                lineheight=1.35,
            )
        document.set_metadata(
            {
                "title": f"Synthetic LegalBridge {source.title}",
                "author": "LegalBridge India Hackathon",
                "subject": "Fictional synthetic demonstration data",
            }
        )
        document.save(path, garbage=4, deflate=True)
    finally:
        document.close()


def _write_docx(path: Path, source: SourceSpec, case: CaseSpec) -> None:
    document = Document()
    document.add_heading(f"Synthetic {source.title}", level=0)
    document.add_paragraph(SYNTHETIC_NOTICE)
    for heading, body in (
        (
            "Demonstration identifiers",
            f"Case {case.number}; category {source.category}; synthetic record only.",
        ),
        (
            "Controlled narrative",
            "A fictional staff member recorded a handling step and timestamp for "
            "demonstration. The entry is not evidence of wrongdoing.",
        ),
        (
            "Review observations",
            "The source text requires attorney verification and creates no legal finding.",
        ),
        (
            "Safety limitation",
            "No real person, client, official record, password, token, or secret is present.",
        ),
    ):
        document.add_heading(heading, level=1)
        document.add_paragraph(body)
    document.core_properties.title = f"Synthetic LegalBridge {source.title}"
    document.core_properties.subject = "Fictional synthetic demonstration data"
    document.save(path)


def _write_txt(path: Path, source: SourceSpec, case: CaseSpec) -> None:
    sections = [
        (
            f"SYNTHETIC {source.title.upper()} - SECTION 1\n"
            f"{SYNTHETIC_NOTICE}\n\nCase: {case.number}\nCategory: {source.category}\n"
        ),
        (
            f"SYNTHETIC {source.title.upper()} - SECTION 2\n"
            "A controlled fictional entry records a demonstration timestamp and handling "
            "step. It requires attorney verification and is not an official record.\n"
        ),
        (
            f"SYNTHETIC {source.title.upper()} - SECTION 3\n"
            "No legal conclusion, authority, or unsupported claim is asserted. "
            "This file exists only to verify private storage and extraction.\n"
        ),
    ]
    path.write_text("\f".join(sections), encoding="utf-8")


def _generate_source(path: Path, source: SourceSpec, case: CaseSpec) -> None:
    if path.suffix == ".pdf":
        _write_pdf(path, source, case)
    elif path.suffix == ".docx":
        _write_docx(path, source, case)
    else:
        _write_txt(path, source, case)


async def _event_exists(
    session: AsyncSession,
    *,
    organization_id: str,
    event_type: str,
    entity_id: str,
) -> bool:
    return (
        await session.scalar(
            select(AuditEvent.id).where(
                AuditEvent.organization_id == organization_id,
                AuditEvent.event_type == event_type,
                AuditEvent.entity_id == entity_id,
            )
        )
        is not None
    )


async def _provision_workspace(
    session: AsyncSession,
) -> tuple[Organization, User, list[User]]:
    organization = await session.scalar(
        select(Organization).where(Organization.slug == MAIN_ORGANIZATION_SLUG)
    )
    if organization is None:
        organization = Organization(
            name=MAIN_ORGANIZATION_NAME,
            slug=MAIN_ORGANIZATION_SLUG,
            is_active=True,
        )
        session.add(organization)
        await session.flush()
    else:
        organization.name = MAIN_ORGANIZATION_NAME
        organization.is_active = True

    user_specs = (
        StaffSpec(PRIMARY_FULL_NAME, PRIMARY_EMAIL, UserRole.ADMIN),
        *STAFF,
    )
    users: list[User] = []
    for spec in user_specs:
        user = await session.scalar(
            select(User).where(
                User.organization_id == organization.id,
                User.email == spec.email,
            )
        )
        created = user is None
        if user is None:
            initial_password = (
                PRIMARY_PASSWORD if spec.email == PRIMARY_EMAIL else secrets.token_urlsafe(48)
            )
            user = User(
                organization_id=organization.id,
                email=spec.email,
                full_name=spec.full_name,
                password_hash=hash_password(initial_password),
                role=spec.role,
                is_active=True,
            )
            session.add(user)
            await session.flush()
        else:
            user.full_name = spec.full_name
            user.role = spec.role
            user.is_active = True

        if spec.email == PRIMARY_EMAIL:
            try:
                password_matches = verify_password(PRIMARY_PASSWORD, user.password_hash)
            except Exception:
                password_matches = False
            if not password_matches:
                user.password_hash = hash_password(PRIMARY_PASSWORD)
                user.token_version += 1
                await session.execute(
                    update(AuthSession)
                    .where(
                        AuthSession.user_id == user.id,
                        AuthSession.revoked_at.is_(None),
                    )
                    .values(revoked_at=utc_now(), updated_at=utc_now())
                )

        if created or not await _event_exists(
            session,
            organization_id=organization.id,
            event_type="user_provisioned",
            entity_id=user.id,
        ):
            add_audit_event(
                session,
                organization_id=organization.id,
                actor_user_id=None,
                event_type="user_provisioned",
                message="Synthetic jury workspace staff user provisioned.",
                entity_type="user",
                entity_id=user.id,
                metadata={"role": user.role.value, "synthetic": True},
            )
        users.append(user)

    if not await _event_exists(
        session,
        organization_id=organization.id,
        event_type="workspace_bootstrapped",
        entity_id=organization.id,
    ):
        add_audit_event(
            session,
            organization_id=organization.id,
            actor_user_id=users[0].id,
            event_type="workspace_bootstrapped",
            message="LegalBridge synthetic jury workspace bootstrapped.",
            entity_type="organization",
            entity_id=organization.id,
            metadata={"workspace_slug": MAIN_ORGANIZATION_SLUG, "synthetic": True},
        )
    await session.commit()
    return organization, users[0], users[1:]


async def _provision_cases(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    attorneys: list[User],
) -> list[LegalCase]:
    legal_cases: list[LegalCase] = []
    for index, spec in enumerate(CASES):
        legal_case = await session.scalar(
            select(LegalCase).where(
                LegalCase.organization_id == organization.id,
                LegalCase.case_number == spec.number,
            )
        )
        assigned = attorneys[index % len(attorneys)]
        if legal_case is None:
            legal_case = LegalCase(
                organization_id=organization.id,
                case_number=spec.number,
                title=spec.title,
                description=(
                    f"{SYNTHETIC_NOTICE} This database-backed matter demonstrates "
                    f"{spec.allegation_type.lower()} review without real client data."
                ),
                court_name=f"Synthetic District Review Forum {index % 5 + 1}",
                jurisdiction="Closed synthetic Indian demonstration jurisdiction",
                allegation_type=spec.allegation_type,
                status=spec.status,
                created_by_id=primary.id,
                assigned_attorney_id=assigned.id,
            )
            session.add(legal_case)
            await session.flush()
        else:
            legal_case.title = spec.title
            legal_case.description = (
                f"{SYNTHETIC_NOTICE} This database-backed matter demonstrates "
                f"{spec.allegation_type.lower()} review without real client data."
            )
            legal_case.court_name = f"Synthetic District Review Forum {index % 5 + 1}"
            legal_case.jurisdiction = "Closed synthetic Indian demonstration jurisdiction"
            legal_case.allegation_type = spec.allegation_type
            legal_case.status = spec.status
            legal_case.assigned_attorney_id = assigned.id

        for event_type, message in (
            ("case_created", "Synthetic jury case created."),
            ("case_assigned", "Synthetic jury case assigned to fictional counsel."),
            ("case_status_changed", "Synthetic jury case status established."),
        ):
            if not await _event_exists(
                session,
                organization_id=organization.id,
                event_type=event_type,
                entity_id=legal_case.id,
            ):
                add_audit_event(
                    session,
                    organization_id=organization.id,
                    actor_user_id=primary.id,
                    event_type=event_type,
                    message=message,
                    entity_type="case",
                    entity_id=legal_case.id,
                    case_id=legal_case.id,
                    metadata={
                        "case_number": spec.number,
                        "status": spec.status.value,
                        "synthetic": True,
                    },
                )
        legal_cases.append(legal_case)
    await session.commit()
    return legal_cases


async def _ingest_source(
    session: AsyncSession,
    *,
    source: SourceSpec,
    source_path: Path,
    organization: Organization,
    legal_case: LegalCase,
    primary: User,
    settings: Settings,
    storage: StorageService,
) -> DocumentRecord:
    existing = await session.scalar(
        select(DocumentRecord).where(
            DocumentRecord.organization_id == organization.id,
            DocumentRecord.case_id == legal_case.id,
            DocumentRecord.original_filename == source.filename,
        )
    )
    if existing is not None:
        return existing

    document_id = new_uuid()
    add_audit_event(
        session,
        organization_id=organization.id,
        actor_user_id=primary.id,
        event_type="document_upload_started",
        message="Synthetic jury source ingestion started.",
        entity_type="document_record",
        entity_id=document_id,
        case_id=legal_case.id,
        metadata={
            "document_id": document_id,
            "category": source.category,
            "declared_content_type": source.content_type,
        },
    )
    await session.commit()
    staged = storage.stage_file(
        source_path,
        filename=source.filename,
        content_type=source.content_type,
    )
    storage_key: str | None = None
    try:
        storage_key = storage.finalize(
            staged,
            organization_id=organization.id,
            case_id=legal_case.id,
            document_id=document_id,
        )
        record = DocumentRecord(
            id=document_id,
            organization_id=organization.id,
            case_id=legal_case.id,
            original_filename=source.filename,
            content_type=source.content_type,
            size_bytes=staged.size_bytes,
            sha256=staged.sha256,
            category=source.category,
            status="metadata_only",
            storage_key=storage_key,
            storage_backend="local_private",
            extraction_status="uploaded",
            original_uploaded_at=utc_now(),
            created_by_id=primary.id,
        )
        session.add(record)
        await session.flush()
        add_audit_event(
            session,
            organization_id=organization.id,
            actor_user_id=primary.id,
            event_type="document_uploaded",
            message="Synthetic jury source stored privately.",
            entity_type="document_record",
            entity_id=record.id,
            case_id=legal_case.id,
            metadata=document_audit_metadata(record),
        )
        await session.commit()
    except Exception:
        await session.rollback()
        if storage_key:
            storage.delete_key(storage_key)
        raise
    finally:
        storage.discard(staged)

    await process_document(
        session,
        record=record,
        storage=storage,
        settings=settings,
        actor_user_id=primary.id,
    )
    return record


async def _provision_documents(
    session: AsyncSession,
    *,
    organization: Organization,
    primary: User,
    legal_cases: list[LegalCase],
    settings: Settings,
) -> None:
    storage = StorageService(settings.storage_root, settings.max_upload_bytes)
    storage.ensure_ready()
    with tempfile.TemporaryDirectory(prefix="legalbridge-main-jury-") as temporary:
        temporary_root = Path(temporary)
        for index, legal_case in enumerate(legal_cases):
            case_spec = CASES[index]
            sources = FLAGSHIP_SOURCES if index == 0 else _supporting_sources(case_spec.number)
            for source in sources:
                existing = await session.scalar(
                    select(DocumentRecord.id).where(
                        DocumentRecord.organization_id == organization.id,
                        DocumentRecord.case_id == legal_case.id,
                        DocumentRecord.original_filename == source.filename,
                    )
                )
                if existing is not None:
                    continue
                source_path = temporary_root / source.filename
                _generate_source(source_path, source, case_spec)
                await _ingest_source(
                    session,
                    source=source,
                    source_path=source_path,
                    organization=organization,
                    legal_case=legal_case,
                    primary=primary,
                    settings=settings,
                    storage=storage,
                )


async def bootstrap_main(database: Database, settings: Settings) -> BootstrapSummary:
    if not settings.database_url.startswith("postgresql+asyncpg://"):
        raise RuntimeError("The legalbridge-main jury bootstrap requires active PostgreSQL.")
    async with database.session_factory() as session:
        organization, primary, supporting = await _provision_workspace(session)
        attorneys = [user for user in supporting if user.role == UserRole.ATTORNEY]
        legal_cases = await _provision_cases(
            session,
            organization=organization,
            primary=primary,
            attorneys=attorneys,
        )
        await _provision_documents(
            session,
            organization=organization,
            primary=primary,
            legal_cases=legal_cases,
            settings=settings,
        )

        case_count = await session.scalar(
            select(func.count(LegalCase.id)).where(LegalCase.organization_id == organization.id)
        )
        document_count = await session.scalar(
            select(func.count(DocumentRecord.id)).where(
                DocumentRecord.organization_id == organization.id
            )
        )
        source_page_count = await session.scalar(
            select(func.count(DocumentPage.id)).where(
                DocumentPage.organization_id == organization.id
            )
        )
        audit_event_count = await session.scalar(
            select(func.count(AuditEvent.id)).where(AuditEvent.organization_id == organization.id)
        )
        flagship_document_count = await session.scalar(
            select(func.count(DocumentRecord.id)).where(
                DocumentRecord.organization_id == organization.id,
                DocumentRecord.case_id == legal_cases[0].id,
            )
        )
        summary = BootstrapSummary(
            organization_id=organization.id,
            primary_user_id=primary.id,
            supporting_user_count=len(supporting),
            case_count=case_count or 0,
            document_count=document_count or 0,
            source_page_count=source_page_count or 0,
            audit_event_count=audit_event_count or 0,
            flagship_document_count=flagship_document_count or 0,
        )
        if (
            summary.case_count < 15
            or summary.document_count < 45
            or summary.source_page_count < 120
            or summary.audit_event_count < 150
            or summary.flagship_document_count < 8
        ):
            raise RuntimeError(f"Jury bootstrap minimums were not met: {summary}")
        return summary


def _masked_database_host(database_url: str) -> str:
    host = make_url(database_url).host or "unknown"
    labels = host.split(".")
    first = labels[0]
    masked_first = f"{first[:3]}***{first[-2:]}" if len(first) > 5 else "***"
    return ".".join((masked_first, *labels[1:]))


async def main() -> None:
    settings = get_settings()
    if not settings.database_url.startswith("postgresql+asyncpg://"):
        raise RuntimeError("Database engine must be PostgreSQL for the jury bootstrap.")
    print("Database engine: PostgreSQL")
    print(f"Database host: {_masked_database_host(settings.database_url)}")
    database = Database(
        settings.database_url,
        echo=settings.sql_echo,
        ssl_mode=settings.database_ssl,
        pool_size=settings.database_pool_size,
        max_overflow=settings.database_max_overflow,
        pool_timeout=settings.database_pool_timeout,
        pool_recycle=settings.database_pool_recycle,
    )
    try:
        summary = await bootstrap_main(database, settings)
    finally:
        await database.dispose()
    print(f"Organisation ID: {summary.organization_id}")
    print(f"Primary user ID: {summary.primary_user_id}")
    print(f"Supporting users: {summary.supporting_user_count}")
    print(f"Cases: {summary.case_count}")
    print(f"Documents: {summary.document_count}")
    print(f"Source pages: {summary.source_page_count}")
    print(f"Audit events: {summary.audit_event_count}")
    print(f"Flagship documents: {summary.flagship_document_count}")


if __name__ == "__main__":
    asyncio.run(main())
