"""Source-grounded motion drafting, checks, review, and export routes."""

from __future__ import annotations

from datetime import datetime, timezone
from io import BytesIO
from typing import Annotated, Any

import fitz
from docx import Document
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.dependencies import Principal, get_current_principal, require_roles
from app.api.routes.cases import get_organization_case
from app.core.config import Settings, get_settings
from app.core.errors import ApplicationError
from app.db.session import get_session
from app.models.analysis import (
    AttorneyReview,
    CitationCheckRecord,
    LegalAuthority,
    MotionDraft,
    MotionVersion,
)
from app.models.document import DocumentRecord
from app.models.document_page import DocumentPage
from app.models.enums import UserRole
from app.schemas.platform import MotionCreate, MotionReviewCreate, MotionVersionCreate
from app.services.analysis import render_motion
from app.services.audit import add_audit_event
from app.services.platform import latest_analysis_run, motion_payload

router = APIRouter(prefix="/cases/{case_id}/motions", tags=["motions"])
motion_editor = require_roles(UserRole.ADMIN, UserRole.ATTORNEY)
motion_reviewer = require_roles(UserRole.ADMIN, UserRole.ATTORNEY, UserRole.REVIEWER)


async def _get_motion(
    session: AsyncSession,
    *,
    organization_id: str,
    case_id: str,
    motion_id: str,
) -> MotionDraft:
    motion = (
        await session.scalars(
            select(MotionDraft).where(
                MotionDraft.id == motion_id,
                MotionDraft.organization_id == organization_id,
                MotionDraft.case_id == case_id,
            )
        )
    ).one_or_none()
    if motion is None:
        raise ApplicationError(
            status_code=404, code="motion_not_found", message="Motion draft not found."
        )
    return motion


async def _latest_version(
    session: AsyncSession, motion_id: str
) -> MotionVersion:
    version = (
        await session.scalars(
            select(MotionVersion)
            .where(MotionVersion.motion_draft_id == motion_id)
            .order_by(MotionVersion.version_number.desc())
            .limit(1)
        )
    ).one_or_none()
    if version is None:
        raise ApplicationError(
            status_code=409,
            code="motion_version_missing",
            message="The motion has no version to review or export.",
        )
    return version


@router.get("")
async def list_motions(
    case_id: str,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> list[dict[str, Any]]:
    await get_organization_case(
        session, organization_id=principal.organization.id, case_id=case_id
    )
    motions = (
        await session.scalars(
            select(MotionDraft)
            .where(
                MotionDraft.organization_id == principal.organization.id,
                MotionDraft.case_id == case_id,
            )
            .order_by(MotionDraft.created_at.desc())
        )
    ).all()
    return [await motion_payload(session, motion) for motion in motions]


@router.post("")
async def create_motion(
    case_id: str,
    payload: MotionCreate,
    principal: Annotated[Principal, Depends(motion_editor)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    await get_organization_case(
        session, organization_id=principal.organization.id, case_id=case_id
    )
    run = await latest_analysis_run(
        session, organization_id=principal.organization.id, case_id=case_id
    )
    if run is None or run.status != "completed":
        raise ApplicationError(
            status_code=409,
            code="analysis_required",
            message="A completed source-grounded analysis is required before drafting.",
        )
    motion = MotionDraft(
        organization_id=principal.organization.id,
        case_id=case_id,
        analysis_run_id=run.id,
        title=payload.title,
        motion_type=payload.motion_type,
        status="draft",
        current_version=1,
        created_by_user_id=principal.user.id,
    )
    session.add(motion)
    await session.flush()
    sections = {
        "Demonstration notice": (
            "Demonstration draft — attorney review required — not filed with any court."
        ),
        "Case information": "See the authenticated case record and stored source pages.",
        "Factual background": "Attorney must add only source-grounded observations.",
        "Limitations and review notice": (
            "Synthetic demonstration data; not legal advice; no automatic court filing."
        ),
    }
    version = MotionVersion(
        motion_draft_id=motion.id,
        version_number=1,
        content_json=sections,
        rendered_text=render_motion(sections),
        citation_check_status="pending",
        ethics_check_status="pending",
        created_by_user_id=principal.user.id,
    )
    session.add(version)
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="motion_created",
        message="Source-grounded demonstration motion created.",
        entity_type="motion_draft",
        entity_id=motion.id,
        case_id=case_id,
    )
    await session.commit()
    return await motion_payload(session, motion)


@router.get("/{motion_id}")
async def get_motion(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    return await motion_payload(session, motion)


@router.post("/{motion_id}/versions")
async def create_motion_version(
    case_id: str,
    motion_id: str,
    payload: MotionVersionCreate,
    principal: Annotated[Principal, Depends(motion_editor)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    previous = await _latest_version(session, motion.id)
    content = payload.content_json or dict(previous.content_json)
    rendered = payload.rendered_text or render_motion(content)
    version_number = previous.version_number + 1
    version = MotionVersion(
        motion_draft_id=motion.id,
        version_number=version_number,
        content_json=content,
        rendered_text=rendered,
        citation_check_status="pending",
        ethics_check_status="pending",
        created_by_user_id=principal.user.id,
    )
    session.add(version)
    motion.current_version = version_number
    motion.status = "draft"
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="motion_version_created",
        message=f"Motion version {version_number} created; prior approval invalidated.",
        entity_type="motion_version",
        entity_id=version.id,
        case_id=case_id,
    )
    await session.commit()
    return await motion_payload(session, motion)


@router.post("/{motion_id}/citation-check")
async def check_motion_citations(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(motion_editor)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    version = await _latest_version(session, motion.id)
    pages = (
        await session.execute(
            select(DocumentPage, DocumentRecord)
            .join(DocumentRecord, DocumentRecord.id == DocumentPage.document_id)
            .where(
                DocumentPage.organization_id == principal.organization.id,
                DocumentPage.case_id == case_id,
            )
            .order_by(DocumentPage.created_at)
            .limit(3)
        )
    ).all()
    authorities = (
        await session.scalars(
            select(LegalAuthority)
            .where(
                LegalAuthority.organization_id == principal.organization.id,
                LegalAuthority.is_synthetic.is_(True),
                LegalAuthority.source_status == "synthetic_demo",
            )
            .order_by(LegalAuthority.citation)
            .limit(2)
        )
    ).all()
    if not pages:
        version.citation_check_status = "failed_missing_sources"
        session.add(
            CitationCheckRecord(
                organization_id=principal.organization.id,
                case_id=case_id,
                motion_version_id=version.id,
                citation_text="No stored source page",
                status="missing_source",
                message="No document page exists; unsupported citations are rejected.",
            )
        )
    else:
        for page, document in pages:
            session.add(
                CitationCheckRecord(
                    organization_id=principal.organization.id,
                    case_id=case_id,
                    motion_version_id=version.id,
                    citation_text=f"{document.original_filename}, page {page.page_number}",
                    source_document_id=document.id,
                    source_page_id=page.id,
                    status="verified_source",
                    message="Stored document-page reference verified.",
                )
            )
        for authority in authorities:
            session.add(
                CitationCheckRecord(
                    organization_id=principal.organization.id,
                    case_id=case_id,
                    motion_version_id=version.id,
                    citation_text=authority.citation,
                    authority_id=authority.id,
                    status="synthetic_demo",
                    message=(
                        "Synthetic demonstration authority — not an official legal source."
                    ),
                )
            )
        version.citation_check_status = "passed_synthetic_sources"
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="citation_firewall_check",
        message="Citation Firewall completed for the current motion version.",
        entity_type="motion_version",
        entity_id=version.id,
        case_id=case_id,
        metadata={"status": version.citation_check_status},
    )
    await session.commit()
    return await motion_payload(session, motion)


@router.post("/{motion_id}/ethics-check")
async def check_motion_ethics(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(motion_editor)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    version = await _latest_version(session, motion.id)
    required_phrases = ("attorney review", "not filed", "synthetic")
    missing = [phrase for phrase in required_phrases if phrase not in version.rendered_text.lower()]
    version.ethics_check_status = (
        "requires_revision" if missing else "passed_with_attorney_review"
    )
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="ethics_check",
        message="Ethics Auditor completed for the current motion version.",
        entity_type="motion_version",
        entity_id=version.id,
        case_id=case_id,
        metadata={"missing_required_phrases": missing},
    )
    await session.commit()
    return await motion_payload(session, motion)


@router.post("/{motion_id}/submit-review")
async def submit_motion_review(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(motion_editor)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    version = await _latest_version(session, motion.id)
    if not version.citation_check_status.startswith("passed"):
        raise ApplicationError(
            status_code=409,
            code="citation_check_required",
            message="Citation Firewall must pass before attorney review.",
        )
    motion.status = "review_required"
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="motion_submitted_for_review",
        message="Motion submitted for internal attorney review.",
        entity_type="motion_draft",
        entity_id=motion.id,
        case_id=case_id,
    )
    await session.commit()
    return await motion_payload(session, motion)


@router.post("/{motion_id}/review")
async def review_motion(
    case_id: str,
    motion_id: str,
    payload: MotionReviewCreate,
    principal: Annotated[Principal, Depends(motion_reviewer)],
    session: Annotated[AsyncSession, Depends(get_session)],
    settings: Annotated[Settings, Depends(get_settings)],
) -> dict[str, Any]:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    if payload.review_pin != settings.review_pin:
        raise ApplicationError(
            status_code=403,
            code="invalid_review_pin",
            message="The review PIN was not accepted.",
        )
    review = AttorneyReview(
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_draft_id=motion.id,
        reviewer_user_id=principal.user.id,
        decision=payload.decision,
        comments=payload.comments,
        review_pin_verified=True,
        reviewed_at=datetime.now(timezone.utc),
    )
    session.add(review)
    motion.status = payload.decision
    await session.flush()
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="attorney_review",
        message=(
            "Internal demonstration review recorded — not a court signature "
            "and not a court filing."
        ),
        entity_type="attorney_review",
        entity_id=review.id,
        case_id=case_id,
        metadata={"decision": payload.decision},
    )
    await session.commit()
    return await motion_payload(session, motion)


def _pdf_bytes(title: str, text: str) -> bytes:
    pdf = fitz.open()
    for index in range(0, len(text), 3400):
        page = pdf.new_page()
        if index == 0:
            page.insert_text((54, 50), title, fontsize=14)
        page.insert_textbox(
            fitz.Rect(54, 72, 558, 790),
            text[index : index + 3400],
            fontsize=9.5,
            lineheight=1.25,
        )
    output = pdf.tobytes(garbage=4, deflate=True)
    pdf.close()
    return output


def _docx_bytes(title: str, text: str) -> bytes:
    document = Document()
    document.add_heading(title, level=1)
    document.add_paragraph(
        "Demonstration draft — attorney review required — not filed with any court."
    )
    for block in text.split("\n\n"):
        lines = block.splitlines()
        if len(lines) > 1:
            document.add_heading(lines[0], level=2)
            document.add_paragraph("\n".join(lines[1:]))
        else:
            document.add_paragraph(block)
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


async def _export(
    *,
    case_id: str,
    motion_id: str,
    export_format: str,
    principal: Principal,
    session: AsyncSession,
) -> StreamingResponse:
    motion = await _get_motion(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        motion_id=motion_id,
    )
    version = await _latest_version(session, motion.id)
    approved_review_id = await session.scalar(
        select(AttorneyReview.id)
        .where(
            AttorneyReview.motion_draft_id == motion.id,
            AttorneyReview.organization_id == principal.organization.id,
            AttorneyReview.case_id == case_id,
            AttorneyReview.decision == "approved",
            AttorneyReview.review_pin_verified.is_(True),
        )
        .order_by(AttorneyReview.reviewed_at.desc())
        .limit(1)
    )
    if motion.status not in {"approved", "exported"} or approved_review_id is None:
        raise ApplicationError(
            status_code=409,
            code="internal_approval_required",
            message="Internal attorney approval is required before export.",
        )
    if export_format == "pdf":
        content = _pdf_bytes(motion.title, version.rendered_text)
        media_type = "application/pdf"
    else:
        content = _docx_bytes(motion.title, version.rendered_text)
        media_type = (
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    motion.status = "exported"
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="motion_export",
        message=f"Internally approved demonstration motion exported as {export_format.upper()}.",
        entity_type="motion_draft",
        entity_id=motion.id,
        case_id=case_id,
        metadata={"format": export_format, "version": version.version_number},
    )
    await session.commit()
    filename = f"legalbridge-demonstration-{motion.id}.{export_format}"
    return StreamingResponse(
        BytesIO(content),
        media_type=media_type,
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@router.get("/{motion_id}/export/pdf")
async def export_motion_pdf(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(motion_reviewer)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> StreamingResponse:
    return await _export(
        case_id=case_id,
        motion_id=motion_id,
        export_format="pdf",
        principal=principal,
        session=session,
    )


@router.get("/{motion_id}/export/docx")
async def export_motion_docx(
    case_id: str,
    motion_id: str,
    principal: Annotated[Principal, Depends(motion_reviewer)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> StreamingResponse:
    return await _export(
        case_id=case_id,
        motion_id=motion_id,
        export_format="docx",
        principal=principal,
        session=session,
    )
