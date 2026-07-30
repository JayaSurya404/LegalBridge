"""Persisted, case-aware, source-grounded Legal Copilot routes."""

from datetime import date
from io import BytesIO
from typing import Annotated, Any, Literal

import pymupdf
from docx import Document
from fastapi import APIRouter, Depends
from fastapi.responses import StreamingResponse
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.api.dependencies import Principal, get_current_principal
from app.core.config import get_settings
from app.api.routes.cases import get_organization_case
from app.core.errors import ApplicationError
from app.db.session import get_session
from app.models.analysis import CopilotMessage, CopilotThread
from app.schemas.platform import CopilotMessageCreate, CopilotThreadCreate
from app.services.audit import add_audit_event
from app.services.platform import copilot_answer, serialize_model

router = APIRouter(prefix="/cases/{case_id}/copilot", tags=["copilot"])


def _report_bytes(title: str, content: str, references: list[dict[str, str]], format: str) -> tuple[bytes, str]:
    manifest = "\n".join(f"- {item['label']}" for item in references) or "- No source pages available"
    body = f"{title}\nGenerated: {date.today().isoformat()}\n\nEvidence and analysis\n{content}\n\nSource manifest\n{manifest}\n\nAttorney review required. Not legal advice; no automatic court filing."
    if format == "docx":
        document = Document()
        document.add_heading(title, level=0)
        for heading, text in (("Evidence and analysis", content), ("Source manifest", manifest)):
            document.add_heading(heading, level=1)
            document.add_paragraph(text)
        stream = BytesIO(); document.save(stream)
        return stream.getvalue(), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    pdf = pymupdf.open(); page = pdf.new_page(); page.insert_textbox(pymupdf.Rect(48, 48, 547, 795), body, fontsize=10)
    data = pdf.tobytes(); pdf.close()
    return data, "application/pdf"


async def _thread(
    session: AsyncSession,
    *,
    organization_id: str,
    case_id: str,
    thread_id: str,
) -> CopilotThread:
    thread = (
        await session.scalars(
            select(CopilotThread).where(
                CopilotThread.id == thread_id,
                CopilotThread.organization_id == organization_id,
                CopilotThread.case_id == case_id,
            )
        )
    ).one_or_none()
    if thread is None:
        raise ApplicationError(
            status_code=404,
            code="copilot_thread_not_found",
            message="Copilot thread not found.",
        )
    return thread


async def _payload(
    session: AsyncSession, thread: CopilotThread
) -> dict[str, Any]:
    messages = (
        await session.scalars(
            select(CopilotMessage)
            .where(CopilotMessage.thread_id == thread.id)
            .order_by(CopilotMessage.created_at)
        )
    ).all()
    return {
        **serialize_model(thread),
        "messages": [serialize_model(message) for message in messages],
    }


@router.get("/threads")
async def list_threads(
    case_id: str,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> list[dict[str, Any]]:
    await get_organization_case(
        session, organization_id=principal.organization.id, case_id=case_id
    )
    threads = (
        await session.scalars(
            select(CopilotThread)
            .where(
                CopilotThread.organization_id == principal.organization.id,
                CopilotThread.case_id == case_id,
            )
            .order_by(CopilotThread.updated_at.desc())
        )
    ).all()
    return [await _payload(session, thread) for thread in threads]


@router.post("/threads")
async def create_thread(
    case_id: str,
    payload: CopilotThreadCreate,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    await get_organization_case(
        session, organization_id=principal.organization.id, case_id=case_id
    )
    thread = CopilotThread(
        organization_id=principal.organization.id,
        case_id=case_id,
        created_by_user_id=principal.user.id,
        title=payload.title,
    )
    session.add(thread)
    await session.flush()
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="copilot_thread_created",
        message="Case-aware Legal Copilot thread created.",
        entity_type="copilot_thread",
        entity_id=thread.id,
        case_id=case_id,
    )
    await session.commit()
    return await _payload(session, thread)


@router.get("/threads/{thread_id}")
async def get_thread(
    case_id: str,
    thread_id: str,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    thread = await _thread(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        thread_id=thread_id,
    )
    return await _payload(session, thread)


@router.post("/threads/{thread_id}/messages")
async def create_message(
    case_id: str,
    thread_id: str,
    payload: CopilotMessageCreate,
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> dict[str, Any]:
    thread = await _thread(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        thread_id=thread_id,
    )
    user_message = CopilotMessage(
        thread_id=thread.id,
        role="user",
        content=payload.content,
        source_references_json=[],
    )
    session.add(user_message)
    answer, references = await copilot_answer(
        session,
        organization_id=principal.organization.id,
        case_id=case_id,
        question=payload.content,
    )
    assistant_message = CopilotMessage(
        thread_id=thread.id,
        role="assistant",
        content=answer,
        source_references_json=references,
        metadata_json={
            "provider_used": (
                "nvidia_nim"
                if get_settings().ai_provider == "nvidia_nim"
                and not answer.startswith("The retrieved case records contain")
                else "extractive_fallback"
            )
        },
    )
    session.add(assistant_message)
    await session.flush()
    add_audit_event(
        session,
        organization_id=principal.organization.id,
        actor_user_id=principal.user.id,
        event_type="copilot_response",
        message="Legal Copilot produced a deterministic source-grounded response.",
        entity_type="copilot_message",
        entity_id=assistant_message.id,
        case_id=case_id,
        metadata={
            "source_reference_count": len(references),
            "provider_used": assistant_message.metadata_json["provider_used"],
        },
    )
    await session.commit()
    return {
        "user_message": serialize_model(user_message),
        "assistant_message": serialize_model(assistant_message),
    }


@router.get("/threads/{thread_id}/reports/{format}/{kind}")
async def download_report(
    case_id: str,
    thread_id: str,
    format: Literal["pdf", "docx"],
    kind: Literal["chronology", "contradictions", "summary"],
    principal: Annotated[Principal, Depends(get_current_principal)],
    session: Annotated[AsyncSession, Depends(get_session)],
) -> StreamingResponse:
    await _thread(session, organization_id=principal.organization.id, case_id=case_id, thread_id=thread_id)
    question = {"chronology": "Build the complete chronology with sources.", "contradictions": "Explain all contradictions with sources.", "summary": "Summarise the complete case with sources."}[kind]
    answer, references = await copilot_answer(session, organization_id=principal.organization.id, case_id=case_id, question=question)
    content, media_type = _report_bytes(f"LegalBridge {kind.title()} Report", answer, references, format)
    add_audit_event(session, organization_id=principal.organization.id, actor_user_id=principal.user.id, event_type="copilot_report_generated", message=f"Generated {format.upper()} Copilot {kind} report.", entity_type="copilot_thread", entity_id=thread_id, case_id=case_id)
    await session.commit()
    return StreamingResponse(BytesIO(content), media_type=media_type, headers={"Content-Disposition": f'attachment; filename="legalbridge-{kind}.{format}"'})
