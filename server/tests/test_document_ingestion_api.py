"""Focused Phase 5–6 storage, extraction, isolation, and bootstrap tests."""

from __future__ import annotations

import asyncio
import hashlib
import sqlite3
from io import BytesIO
from pathlib import Path

import pymupdf
from docx import Document
from sqlalchemy import func, select

from app.core.config import Settings
from app.db.base import Base
from app.db.session import Database
from app.models.audit import AuditEvent
from app.models.document import DocumentRecord
from app.models.document_page import DocumentPage
from app.scripts.bootstrap_demo import bootstrap_demo
from app.scripts.bootstrap_demo_documents import bootstrap_demo_documents
from tests.conftest import ApiContext


def _create_case(context: ApiContext, number: str = "INGEST-001") -> dict[str, object]:
    response = context.client.post(
        "/api/v1/cases",
        headers=context.access_headers(context.attorney_email),
        json={
            "case_number": number,
            "title": "Synthetic ingestion test matter",
            "description": "Fictional test data only.",
            "status": "active",
            "assigned_attorney_id": context.attorney_id,
        },
    )
    assert response.status_code == 201
    return response.json()


def _upload(
    context: ApiContext,
    *,
    case_id: str,
    filename: str,
    content: bytes,
    content_type: str,
    email: str | None = None,
):
    return context.client.post(
        f"/api/v1/cases/{case_id}/documents/upload",
        headers=context.access_headers(email or context.attorney_email),
        files={"file": (filename, content, content_type)},
        data={"category": "synthetic evidence"},
    )


def _docx_bytes() -> bytes:
    document = Document()
    document.add_heading("Synthetic report", level=1)
    document.add_paragraph(
        "Synthetic hackathon data only. This is not an official record and contains no real person."
    )
    table = document.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Field"
    table.rows[0].cells[1].text = "Value"
    table.rows[1].cells[0].text = "Fictional ID"
    table.rows[1].cells[1].text = "FIC-001"
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _pdf_bytes(*, blank: bool = False, pages: int = 2) -> bytes:
    document = pymupdf.open()
    try:
        for index in range(pages):
            page = document.new_page()
            if not blank:
                page.insert_text(
                    (72, 72),
                    (
                        f"Synthetic PDF physical page {index + 1}. "
                        "This fictional source is not an official record and contains enough text."
                    ),
                    fontsize=11,
                )
        return document.tobytes(garbage=4, deflate=True)
    finally:
        document.close()


def test_valid_txt_upload_server_hash_pages_download_reprocess_and_delete(
    context: ApiContext,
) -> None:
    legal_case = _create_case(context)
    case_id = str(legal_case["id"])
    content = (
        b"Synthetic section one. Not an official record.\n\n"
        b"Fictional identifier TXT-001."
        b"\f"
        b"Synthetic section two with reviewable extracted source text."
    )
    response = _upload(
        context,
        case_id=case_id,
        filename="synthetic-memo.txt",
        content=content,
        content_type="text/plain",
    )

    assert response.status_code == 201
    created = response.json()
    assert created["sha256"] == hashlib.sha256(content).hexdigest()
    assert created["extraction_status"] == "processed"
    assert created["page_count"] == 2
    assert created["extracted_character_count"] > 0
    assert created["binary_exists"] is True
    assert [page["page_number"] for page in created["pages"]] == [1, 2]
    assert all("Text page" in page["page_label"] for page in created["pages"])
    assert "Synthetic section two" in created["pages"][1]["extracted_text"]

    detail = context.client.get(
        f"/api/v1/cases/{case_id}/documents/{created['id']}",
        headers=context.access_headers(context.reviewer_email),
    )
    assert detail.status_code == 200
    assert detail.json()["pages"] == created["pages"]

    downloaded = context.client.get(
        f"/api/v1/cases/{case_id}/documents/{created['id']}/download",
        headers=context.access_headers(context.reviewer_email),
    )
    assert downloaded.status_code == 200
    assert downloaded.content == content
    assert "attachment" in downloaded.headers["content-disposition"]

    duplicate = _upload(
        context,
        case_id=case_id,
        filename="duplicate.txt",
        content=content,
        content_type="text/plain",
    )
    assert duplicate.status_code == 409
    assert duplicate.json()["error"]["code"] == "duplicate_document_sha256"

    reprocessed = context.client.post(
        f"/api/v1/cases/{case_id}/documents/{created['id']}/reprocess",
        headers=context.access_headers(context.attorney_email),
    )
    assert reprocessed.status_code == 200
    assert reprocessed.json()["page_count"] == 2

    stored_path = (
        context.storage_root
        / str(created["organization_id"])
        / case_id
        / str(created["id"])
        / "original.txt"
    )
    assert stored_path.read_bytes() == content
    deleted = context.client.delete(
        f"/api/v1/cases/{case_id}/documents/{created['id']}",
        headers=context.access_headers(context.attorney_email),
    )
    assert deleted.status_code == 204
    assert not stored_path.exists()
    with sqlite3.connect(context.database_path) as connection:
        page_count = connection.execute(
            "SELECT count(*) FROM document_pages WHERE document_id = ?",
            (created["id"],),
        ).fetchone()[0]
    assert page_count == 0

    event_types = [
        event["event_type"]
        for event in context.client.get(
            f"/api/v1/cases/{case_id}/audit-events",
            headers=context.access_headers(context.attorney_email),
        ).json()
    ]
    for expected in (
        "document_upload_started",
        "document_uploaded",
        "document_extraction_started",
        "document_extracted",
        "document_downloaded",
        "document_reprocessed",
        "document_deleted",
        "document_validation_failed",
    ):
        assert expected in event_types


def test_valid_docx_upload_extracts_ordered_logical_sections(
    context: ApiContext,
) -> None:
    legal_case = _create_case(context)
    response = _upload(
        context,
        case_id=str(legal_case["id"]),
        filename="synthetic-report.docx",
        content=_docx_bytes(),
        content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["extraction_status"] == "processed"
    assert payload["parser_name"] == "python-docx"
    assert payload["page_count"] >= 1
    assert payload["pages"][0]["page_label"].startswith("Logical DOCX section")
    assert "Synthetic report" in payload["pages"][0]["extracted_text"]
    assert "Fictional ID | FIC-001" in payload["pages"][0]["extracted_text"]


def test_valid_pdf_upload_preserves_physical_pages(context: ApiContext) -> None:
    legal_case = _create_case(context)
    response = _upload(
        context,
        case_id=str(legal_case["id"]),
        filename="synthetic-transcript.pdf",
        content=_pdf_bytes(pages=2),
        content_type="application/pdf",
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["extraction_status"] == "processed"
    assert payload["parser_name"] == "PyMuPDF"
    assert payload["page_count"] == 2
    assert [page["page_label"] for page in payload["pages"]] == [
        "Physical PDF page 1",
        "Physical PDF page 2",
    ]


def test_image_only_pdf_truthfully_requires_ocr_when_disabled(
    context: ApiContext,
) -> None:
    legal_case = _create_case(context)
    response = _upload(
        context,
        case_id=str(legal_case["id"]),
        filename="synthetic-scan.pdf",
        content=_pdf_bytes(blank=True, pages=1),
        content_type="application/pdf",
    )

    assert response.status_code == 201
    payload = response.json()
    assert payload["extraction_status"] == "ocr_required"
    assert payload["page_count"] == 1
    assert payload["pages"][0]["extracted_text"] == ""
    assert payload["pages"][0]["extraction_method"] == "ocr_required"
    assert "OCR is unavailable or disabled" in payload["extraction_error"]


def test_invalid_uploads_return_controlled_statuses(context: ApiContext) -> None:
    legal_case = _create_case(context)
    case_id = str(legal_case["id"])
    cases = [
        ("bad.exe", b"not executable", "application/octet-stream", 400, "unsupported_extension"),
        ("mismatch.txt", b"plain text", "application/pdf", 400, "content_type_mismatch"),
        ("bad.pdf", b"not a pdf", "application/pdf", 400, "invalid_pdf_signature"),
        (
            "bad.docx",
            b"PK\x03\x04not-an-office-container",
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            400,
            "invalid_docx_container",
        ),
        ("empty.txt", b"", "text/plain", 400, "empty_file"),
        ("binary.txt", b"\x00\x01\x02\x00" * 100, "text/plain", 400, "binary_text_file"),
    ]
    for filename, content, content_type, expected_status, expected_code in cases:
        response = _upload(
            context,
            case_id=case_id,
            filename=filename,
            content=content,
            content_type=content_type,
        )
        assert response.status_code == expected_status
        assert response.json()["error"]["code"] == expected_code

    oversized = _upload(
        context,
        case_id=case_id,
        filename="oversized.txt",
        content=b"a" * (1024 * 1024 + 1),
        content_type="text/plain",
    )
    assert oversized.status_code == 413
    assert oversized.json()["error"]["code"] == "upload_too_large"


def test_reviewer_is_read_only_and_organization_isolation_is_404(
    context: ApiContext,
) -> None:
    legal_case = _create_case(context)
    case_id = str(legal_case["id"])
    created = _upload(
        context,
        case_id=case_id,
        filename="permissions.txt",
        content=b"Synthetic permissions document with enough extracted text for review.",
        content_type="text/plain",
    ).json()
    reviewer_headers = context.access_headers(context.reviewer_email)

    assert (
        context.client.get(
            f"/api/v1/cases/{case_id}/documents",
            headers=reviewer_headers,
        ).status_code
        == 200
    )
    assert (
        context.client.get(
            f"/api/v1/cases/{case_id}/documents/{created['id']}",
            headers=reviewer_headers,
        ).status_code
        == 200
    )
    assert (
        context.client.get(
            f"/api/v1/cases/{case_id}/documents/{created['id']}/download",
            headers=reviewer_headers,
        ).status_code
        == 200
    )
    assert (
        _upload(
            context,
            case_id=case_id,
            filename="reviewer.txt",
            content=b"Reviewer cannot upload this synthetic file.",
            content_type="text/plain",
            email=context.reviewer_email,
        ).status_code
        == 403
    )
    assert (
        context.client.post(
            f"/api/v1/cases/{case_id}/documents/{created['id']}/reprocess",
            headers=reviewer_headers,
        ).status_code
        == 403
    )
    assert (
        context.client.delete(
            f"/api/v1/cases/{case_id}/documents/{created['id']}",
            headers=reviewer_headers,
        ).status_code
        == 403
    )

    other_headers = context.access_headers(
        context.other_admin_email,
        organization_slug="other-legal-aid",
    )
    for path in (
        f"/api/v1/cases/{case_id}/documents",
        f"/api/v1/cases/{case_id}/documents/{created['id']}",
        f"/api/v1/cases/{case_id}/documents/{created['id']}/download",
    ):
        assert context.client.get(path, headers=other_headers).status_code == 404


async def _bootstrap_twice(settings: Settings) -> tuple[int, int, int, int, int]:
    database = Database(settings.database_url)
    try:
        async with database.engine.begin() as connection:
            await connection.run_sync(Base.metadata.create_all)
        await bootstrap_demo(database)
        first = await bootstrap_demo_documents(database, settings)
        second = await bootstrap_demo_documents(database, settings)
        async with database.session_factory() as session:
            document_count = await session.scalar(select(func.count(DocumentRecord.id)))
            page_count = await session.scalar(select(func.count(DocumentPage.id)))
            event_count = await session.scalar(
                select(func.count(AuditEvent.id)).where(
                    AuditEvent.event_type == "demo_documents_bootstrapped"
                )
            )
        return first, second, int(document_count or 0), int(page_count or 0), int(event_count or 0)
    finally:
        await database.dispose()


def test_demo_document_bootstrap_is_idempotent(tmp_path: Path) -> None:
    database_path = (tmp_path / "bootstrap.db").as_posix()
    settings = Settings(
        environment="test",
        database_url=f"sqlite+aiosqlite:///{database_path}",
        storage_root=tmp_path / "demo-uploads",
        jwt_secret="test-only-jwt-secret-that-is-longer-than-thirty-two-characters",
        ocr_enabled=False,
    )

    first, second, document_count, page_count, event_count = asyncio.run(_bootstrap_twice(settings))
    assert first == second == document_count == 3
    assert page_count >= 6
    assert event_count == 1
